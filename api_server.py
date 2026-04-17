#!/usr/bin/env python3
"""
Flask API Backend for Requirements Engineering System
====================================================

This module provides REST API endpoints for the React frontend to interact
with the requirements processing system.
"""

from pathlib import Path
from typing import Any, Optional

# Load .env from project root before anything reads REPLICATE_API_TOKEN / secrets
try:
    from dotenv import load_dotenv

    load_dotenv(Path(__file__).resolve().parent / ".env")
except ImportError:
    pass

from flask import Flask, request, jsonify, send_file, send_from_directory, Response, stream_with_context
from flask_cors import CORS
from werkzeug.exceptions import RequestEntityTooLarge, BadRequest
import json
import os
import tempfile
import time
from datetime import datetime
import logging
import re
import binascii
from urllib.parse import unquote_plus
from collections import defaultdict, deque
import replicate

# Import our existing modules
from main_orchestrator import RequirementsOrchestrator
from srs_generator import SRSGenerator
from generation.srs_generator import RAGSRSGenerator
from json_to_srs_pdf import load_srs_from_json, render_html, save_pdf_or_html
from srs_model_generator import SRSModelGenerator
from generation.textual_usecase_generator import TextualUseCaseGenerator
from generation.usecase_diagram_generator import UseCaseDiagramGenerator
from input_processing.ambiguity_detection import AmbiguityDetector
from input_processing.requirement_refinement import RequirementRefiner
from evaluation.manual_metrics_engine import ManualMetricsEngine
from evaluation.conflict_metric import ConflictMetric
from evaluation.nfr_specificity_metric import NFRSpecificityMetric
from evaluation.professional_style_metric import ProfessionalStyleMetric
import base64
import threading
import uuid
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# CRA emits assets under frontend/build/static/; disable Flask's default /static -> ./static
# (would otherwise serve empty /app/static from Dockerfile and return 404 for JS/CSS).
FRONTEND_BUILD_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "frontend", "build")
app = Flask(__name__, static_folder=None)

# CORS configuration
# In production, restrict to specific origins for security
allowed_origins = os.environ.get('ALLOWED_ORIGINS', '*').split(',')
_cors_expose = ['Content-Disposition', 'Content-Type', 'X-Export-Format']
if allowed_origins == ['*']:
    CORS(app, expose_headers=_cors_expose)  # Allow all origins (development)
else:
    CORS(app, origins=allowed_origins, expose_headers=_cors_expose)

try:
    _max_mb = max(1, min(500, int(os.environ.get("MAX_CONTENT_MB", "80"))))
except (TypeError, ValueError):
    _max_mb = 80
app.config["MAX_CONTENT_LENGTH"] = _max_mb * 1024 * 1024

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Lightweight in-memory guardrails (no external dependency).
_SECURITY_WINDOW_SECONDS = int(os.environ.get("SECURITY_WINDOW_SECONDS", "300"))
_SECURITY_RATE_LIMIT = int(os.environ.get("SECURITY_RATE_LIMIT", "80"))
_SECURITY_RATE_LIMIT_HEAVY = int(os.environ.get("SECURITY_RATE_LIMIT_HEAVY", "20"))
_SECURITY_ATTACK_BLOCK_SECONDS = int(os.environ.get("SECURITY_ATTACK_BLOCK_SECONDS", "600"))
_SECURITY_MAX_EVENTS_PER_IP = int(os.environ.get("SECURITY_MAX_EVENTS_PER_IP", "200"))
_SECURITY_LOCK = threading.Lock()
_REQUEST_EVENTS: dict[str, deque[float]] = defaultdict(deque)
_SUSPICIOUS_EVENTS: dict[str, deque[float]] = defaultdict(deque)
_BLOCKED_UNTIL: dict[str, float] = {}
_HEAVY_ROUTES = frozenset(
    {
        "/api/process-and-generate-srs",
        "/api/process-audio",
        "/api/process-batch",
        "/api/generate-srs",
        "/api/generate-srs-stream",
        "/api/generate-srs-compare",
        "/api/clarification-copilot",
        "/api/clarification-copilot-turn",
    }
)


@app.errorhandler(RequestEntityTooLarge)
def _handle_request_entity_too_large(e):
    if request.path.startswith("/api/"):
        return jsonify(
            {
                "error": (
                    f"Request body exceeds server limit ({_max_mb} MB). "
                    "Use a smaller file, shorter text, or raise MAX_CONTENT_MB in the server environment."
                )
            }
        ), 413
    return e


@app.errorhandler(BadRequest)
def _handle_bad_request(e):
    if request.path.startswith("/api/"):
        desc = getattr(e, "description", None)
        if isinstance(desc, str) and "Failed to decode JSON" in desc:
            return jsonify({"error": "Invalid or malformed JSON body."}), 400
        if isinstance(desc, str) and desc.strip() and desc.lower() != "bad request":
            return jsonify({"error": desc.strip()}), 400
        return jsonify({"error": "Bad request."}), 400
    return e


@app.before_request
def _security_pre_request_guard():
    """
    Per-IP rate limiting and temporary blocking for repeated suspicious activity.
    """
    if not request.path.startswith("/api/"):
        return None
    ip = _get_client_ip()
    now = time.time()
    limit = _SECURITY_RATE_LIMIT_HEAVY if request.path in _HEAVY_ROUTES else _SECURITY_RATE_LIMIT

    with _SECURITY_LOCK:
        blocked_until = _BLOCKED_UNTIL.get(ip, 0.0)
        if blocked_until > now:
            retry = int(max(1, blocked_until - now))
            return jsonify(
                {
                    "error": "Too many suspicious requests from this client. Try again later.",
                    "retry_after_seconds": retry,
                }
            ), 429

        events = _REQUEST_EVENTS[ip]
        events.append(now)
        while events and (now - events[0]) > _SECURITY_WINDOW_SECONDS:
            events.popleft()
        if len(events) > _SECURITY_MAX_EVENTS_PER_IP:
            while len(events) > _SECURITY_MAX_EVENTS_PER_IP:
                events.popleft()

        if len(events) > limit:
            _BLOCKED_UNTIL[ip] = now + min(_SECURITY_ATTACK_BLOCK_SECONDS, 180)
            logger.warning("SECURITY rate_limit ip=%s path=%s count=%s", ip, request.path, len(events))
            return jsonify(
                {
                    "error": "Rate limit exceeded. Please slow down and retry.",
                    "retry_after_seconds": 120,
                }
            ), 429
    return None


# --- Human expert review queue (JSON file; shared across clients on same server) ---
_expert_reviews_lock = threading.Lock()
EXPERT_REVIEWS_PATH = Path(__file__).resolve().parent / "data" / "expert_reviews.json"


def _ensure_expert_reviews_dir():
    EXPERT_REVIEWS_PATH.parent.mkdir(parents=True, exist_ok=True)


def _load_expert_reviews() -> list:
    _ensure_expert_reviews_dir()
    if not EXPERT_REVIEWS_PATH.exists():
        return []
    try:
        with open(EXPERT_REVIEWS_PATH, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, list) else []
    except Exception as e:
        logger.warning("Could not load expert reviews file: %s", e)
        return []


def _save_expert_reviews(reviews: list) -> bool:
    try:
        _ensure_expert_reviews_dir()
        tmp_path = EXPERT_REVIEWS_PATH.with_suffix(".json.tmp")
        with open(tmp_path, "w", encoding="utf-8") as f:
            json.dump(reviews, f, ensure_ascii=False, indent=2)
        tmp_path.replace(EXPERT_REVIEWS_PATH)
        return True
    except OSError as e:
        logger.error("Could not save expert reviews file: %s", e)
        return False


def _find_review_by_id(reviews: list, rid: str) -> int | None:
    for i, r in enumerate(reviews):
        if isinstance(r, dict) and r.get("id") == rid:
            return i
    return None


def _normalize_review_entry(r: dict) -> dict:
    """Ensure chat_messages exists for API responses and in-memory updates."""
    if not isinstance(r, dict):
        return r
    cm = r.get("chat_messages")
    if not isinstance(cm, list):
        return {**r, "chat_messages": []}
    return r


def _compute_kb_quality_metrics(raw_text: str) -> dict:
    """
    Same scoring pipeline as rag/evaluate_srs_kb.py (knowledge-base quality metrics).
    Returns JSON-serializable floats in 0..1 for the generated SRS text.
    """
    if not raw_text or len(raw_text.strip()) < 80:
        return {}
    try:
        from rag.evaluate_srs_kb import KBEvaluator

        raw = KBEvaluator().evaluate_document(raw_text)
    except Exception as e:
        logger.warning("KB quality metrics skipped: %s", e)
        return {}
    out = {}
    for k, v in raw.items():
        try:
            if hasattr(v, "item"):
                v = v.item()
            if isinstance(v, (int, float)):
                fv = float(v)
                if fv != fv:  # NaN
                    continue
                out[k] = round(fv, 4)
            else:
                out[k] = v
        except (TypeError, ValueError):
            continue
    return out


def _build_srs_quality_detail_table(metrics: dict) -> list:
    """
    Human-readable rows for structural / wording quality metrics (same signals as KB evaluation).
    Scores are 0..1 in metrics; UI can show as percent.
    """
    if not isinstance(metrics, dict):
        return []

    def pct(key: str) -> str | None:
        v = metrics.get(key)
        try:
            x = float(v)
            return f"{round(max(0.0, min(1.0, x)) * 100)}%"
        except (TypeError, ValueError):
            return None

    def cnt(key: str) -> int | None:
        v = metrics.get(key)
        try:
            return int(v)
        except (TypeError, ValueError):
            return None

    rows = [
        {
            "key": "arm_overall_score",
            "label": "Overall (structural wording composite)",
            "score_key": "arm_overall_score",
            "score_display": pct("arm_overall_score"),
            "count": None,
            "count_label": None,
            "interpretation": (
                "Single blended score from imperative wording, weak phrases, optional language, "
                "continuance phrases (e.g. 'listed below'), figure/table references, incomplete markers "
                "(TBD, etc.), and vague adjectives. Higher means the SRS reads more like clear, verifiable "
                "requirements."
            ),
        },
        {
            "key": "arm_imperative_quality",
            "label": "Imperative requirements",
            "score_key": "arm_imperative_quality",
            "score_display": pct("arm_imperative_quality"),
            "count": cnt("arm_imperative_count"),
            "count_label": "Lines with shall / must / required to",
            "interpretation": (
                "Share of detected requirement lines that use mandatory language (shall, must, required to). "
                "Strong specifications favor direct obligations; higher is better for testable specs."
            ),
        },
        {
            "key": "arm_weak_phrase_quality",
            "label": "Weak phrases avoided",
            "score_key": "arm_weak_phrase_quality",
            "score_display": pct("arm_weak_phrase_quality"),
            "count": cnt("arm_weak_phrase_count"),
            "count_label": "Weak-phrase hits in document",
            "interpretation": (
                "Penalizes terms like 'adequate', 'as appropriate', 'user-friendly', 'efficient' without "
                "measurable criteria. Lower counts yield a higher score."
            ),
        },
        {
            "key": "arm_optionality_quality",
            "label": "Optionality kept low",
            "score_key": "arm_optionality_quality",
            "score_display": pct("arm_optionality_quality"),
            "count": cnt("arm_option_count"),
            "count_label": "Optional / may / could style hits",
            "interpretation": (
                "Counts optional phrasing (may, can, could, optionally). Requirements should be definitive "
                "unless flexibility is intentional; fewer hits score higher."
            ),
        },
        {
            "key": "arm_continuance_quality",
            "label": "Continuance phrases avoided",
            "score_key": "arm_continuance_quality",
            "score_display": pct("arm_continuance_quality"),
            "count": cnt("arm_continuance_count"),
            "count_label": "'Listed below' / 'as follows' style",
            "interpretation": (
                "Phrases that defer content to another place weaken traceability in the same sentence. "
                "Prefer self-contained statements where possible."
            ),
        },
        {
            "key": "arm_directive_quality",
            "label": "Figure/table directives avoided",
            "score_key": "arm_directive_quality",
            "score_display": pct("arm_directive_quality"),
            "count": cnt("arm_directive_count"),
            "count_label": "Figure / table / note: references",
            "interpretation": (
                "Heavy reliance on 'see Figure X' without restating the requirement can hurt verifiability. "
                "Some references are fine; many in requirement text lower this score."
            ),
        },
        {
            "key": "arm_incomplete_quality",
            "label": "Incomplete markers avoided",
            "score_key": "arm_incomplete_quality",
            "score_display": pct("arm_incomplete_quality"),
            "count": cnt("arm_incomplete_count"),
            "count_label": "TBD / to be determined / placeholder hits",
            "interpretation": (
                "Detects unfinished placeholders. Requirements should be complete enough to "
                "implement and verify; replace TBDs with concrete text where possible."
            ),
        },
        {
            "key": "arm_ambiguity_quality",
            "label": "Ambiguous adjectives avoided",
            "score_key": "arm_ambiguity_quality",
            "score_display": pct("arm_ambiguity_quality"),
            "count": cnt("arm_ambiguity_count"),
            "count_label": "Vague terms (fast, robust, minimal, …)",
            "interpretation": (
                "Flags subjective adjectives that often need numeric or testable criteria. "
                "Fewer such terms yield a higher score."
            ),
        },
    ]
    return [r for r in rows if r["score_display"] is not None]


@app.route('/api/evaluate-srs-kb-metrics', methods=['POST'])
def evaluate_srs_kb_metrics():
    """
    Run KB-style quality scoring on SRS plain text (includes structural wording metrics from the KB pipeline).
    Returns flat `metrics` plus `srs_quality_table` (detail rows for the metrics page).
    """
    try:
        data = request.get_json(silent=True) or {}
        raw_text = sanitize_user_input(str(data.get("raw_text", "") or ""))
        if len(raw_text.strip()) < 80:
            return jsonify(
                {
                    "error": "raw_text is too short (need at least ~80 characters)",
                    "metrics": {},
                    "srs_quality_table": [],
                }
            ), 400
        has_inj, _ = detect_prompt_injection(raw_text)
        if has_inj:
            logger.warning("evaluate-srs-kb-metrics: blocked prompt-injection heuristics")
            return jsonify(
                {
                    "error": "Text content validation failed",
                    "validation_errors": [
                        "Input appears to contain prompt-injection patterns. Paste legitimate SRS text only."
                    ],
                    "security_issue": True,
                    "metrics": {},
                    "srs_quality_table": [],
                }
            ), 403
        metrics = _compute_kb_quality_metrics(raw_text)
        if not metrics:
            return jsonify(
                {
                    "error": "Could not compute metrics for this text.",
                    "metrics": {},
                    "srs_quality_table": [],
                }
            ), 500
        detail_table = _build_srs_quality_detail_table(metrics)
        return jsonify({"metrics": metrics, "srs_quality_table": detail_table})
    except Exception as e:
        logger.error("evaluate-srs-kb-metrics failed: %s", e, exc_info=True)
        return jsonify({"error": str(e), "metrics": {}, "srs_quality_table": []}), 500


def _rtm_tokenize(text: str) -> set:
    txt = str(text or "").lower()
    txt = re.sub(r"[^a-z0-9\s]", " ", txt)
    tokens = {t for t in txt.split() if len(t) >= 3}
    stop = {
        "the", "and", "for", "with", "from", "that", "this", "shall", "should", "must",
        "system", "user", "users", "will", "into", "such", "when", "then", "their",
    }
    return {t for t in tokens if t not in stop}


def _rtm_similarity(a: str, b: str) -> float:
    ta = _rtm_tokenize(a)
    tb = _rtm_tokenize(b)
    if not ta or not tb:
        return 0.0
    inter = len(ta.intersection(tb))
    denom = max(1, len(ta.union(tb)))
    return inter / denom


def _normalize_req_id(value: str) -> str:
    raw = str(value or "").upper().strip()
    m = re.search(r"\b(FR|NFR)\s*[-_ ]?\s*(\d+)\b", raw)
    if not m:
        return ""
    return f"{m.group(1)}-{m.group(2)}"


def _extract_req_ids(text: str) -> set:
    found = set()
    for m in re.finditer(r"\b(FR|NFR)\s*[-_ ]?\s*(\d+)\b", str(text or ""), flags=re.IGNORECASE):
        found.add(f"{m.group(1).upper()}-{m.group(2)}")
    return found


def _canonical_uc_name(name: str) -> str:
    # Strip punctuation and non-semantic prefixes for robust cross-artifact comparison.
    txt = re.sub(r"[^a-z0-9\s]", " ", str(name or "").lower())
    txt = re.sub(r"\b(use\s*case|uc)\s*[-_ ]?\d+\b", " ", txt)
    txt = re.sub(r"\s+", " ", txt).strip()
    return txt


def _extract_requirements_for_rtm(srs_data: dict) -> list:
    rows = []
    sections = (srs_data or {}).get("sections", {}) if isinstance(srs_data, dict) else {}
    sr = sections.get("specific_requirements", {}) if isinstance(sections, dict) else {}
    frs = sr.get("functional_requirements", []) if isinstance(sr, dict) else []
    if isinstance(frs, list):
        for idx, fr in enumerate(frs, start=1):
            req_id = f"FR-{idx}"
            req_text = ""
            if isinstance(fr, dict):
                req_id = str(fr.get("id") or req_id)
                req_text = str(
                    fr.get("description")
                    or " ".join(
                        str(fr.get(k, "")).strip()
                        for k in ("input", "processing", "output")
                        if str(fr.get(k, "")).strip()
                    )
                ).strip()
            else:
                req_text = str(fr).strip()
                m = re.search(r"\b(FR-\d+)\b", req_text, flags=re.IGNORECASE)
                if m:
                    req_id = m.group(1).upper()
            if req_text:
                rows.append({"req_id": req_id.upper(), "type": "FR", "text": req_text})

    perf = sr.get("performance_requirements", {}) if isinstance(sr, dict) else {}
    attrs = sr.get("software_system_attributes", {}) if isinstance(sr, dict) else {}
    nfr_idx = 1
    for block in (perf, attrs):
        if isinstance(block, dict):
            for key, value in block.items():
                txt = str(value or "").strip()
                if not txt:
                    continue
                rows.append(
                    {
                        "req_id": f"NFR-{nfr_idx}",
                        "type": "NFR",
                        "text": f"{str(key).replace('_', ' ').title()}: {txt}",
                    }
                )
                nfr_idx += 1

    raw = str((srs_data or {}).get("raw_text", "") or "")
    if raw:
        for m in re.finditer(r"\b(FR-\d+|NFR-\d+)\b[^\n]*", raw, flags=re.IGNORECASE):
            line = m.group(0).strip()
            rid = m.group(1).upper()
            if not any(r["req_id"] == rid for r in rows):
                rows.append({"req_id": rid, "type": "NFR" if rid.startswith("NFR-") else "FR", "text": line})
    return rows


def _extract_textual_usecases_for_rtm(use_case_data: dict) -> list:
    textual = (use_case_data or {}).get("textual_usecases", {}) if isinstance(use_case_data, dict) else {}
    out = []
    uc_list = textual.get("use_cases", []) if isinstance(textual, dict) else []
    if isinstance(uc_list, list):
        for idx, uc in enumerate(uc_list, start=1):
            if not isinstance(uc, dict):
                continue
            name = str(uc.get("use_case_name") or f"UC-{idx}").strip()
            scenario = str(uc.get("main_success_scenario") or "").strip()
            preconditions = str(uc.get("preconditions") or "").strip()
            full_text = " ".join(part for part in (name, preconditions, scenario) if part).strip()
            ref_ids = set()
            ref_ids.update(_extract_req_ids(uc.get("id", "")))
            ref_ids.update(_extract_req_ids(preconditions))
            ref_ids.update(_extract_req_ids(scenario))
            out.append(
                {
                    "id": f"TUC-{idx}",
                    "name": name,
                    "text": full_text,
                    "req_ids": sorted(ref_ids),
                    "canonical_name": _canonical_uc_name(name),
                }
            )
    if not out:
        rendered = str(textual.get("text", "") if isinstance(textual, dict) else "")
        blocks = [b.strip() for b in re.split(r"\n\s*\n", rendered) if b.strip()]
        for idx, block in enumerate(blocks, start=1):
            m = re.search(r"Use Case Name:\s*(.+)", block, flags=re.IGNORECASE)
            name = (m.group(1).strip() if m else f"UC-{idx}")[:180]
            out.append(
                {
                    "id": f"TUC-{idx}",
                    "name": name,
                    "text": block[:800],
                    "req_ids": sorted(_extract_req_ids(block)),
                    "canonical_name": _canonical_uc_name(name),
                }
            )
    return out


def _extract_diagram_usecases_for_rtm(use_case_data: dict) -> list:
    diagram = (use_case_data or {}).get("diagram", {}) if isinstance(use_case_data, dict) else {}
    puml = str(diagram.get("plantuml_code", "") if isinstance(diagram, dict) else "")
    names = []
    for m in re.finditer(r'usecase\s+"([^"]+)"\s+as\s+([A-Za-z0-9_]+)', puml, flags=re.IGNORECASE):
        name = m.group(1).strip()
        names.append(
            {
                "id": f"DUC-{m.group(2)}",
                "name": name,
                "text": name,
                "req_ids": [],
                "canonical_name": _canonical_uc_name(name),
            }
        )
    if not names:
        for idx, m in enumerate(re.finditer(r"\(([^)]+)\)", puml), start=1):
            nm = m.group(1).strip()
            if nm and not any(n["name"] == nm for n in names):
                names.append(
                    {
                        "id": f"DUC-{idx}",
                        "name": nm,
                        "text": nm,
                        "req_ids": [],
                        "canonical_name": _canonical_uc_name(nm),
                    }
                )
    uc_ref_map = {}
    for m in re.finditer(
        r'([A-Za-z0-9_]+)\s*\.\.\s*note\s*(?:right|left|top|bottom)?\s*:?\s*([^\n]*)',
        puml,
        flags=re.IGNORECASE,
    ):
        alias = str(m.group(1) or "").strip()
        note_txt = str(m.group(2) or "").strip()
        if not alias or not note_txt:
            continue
        uc_ref_map.setdefault(alias, set()).update(_extract_req_ids(note_txt))
    for uc in names:
        alias = uc["id"].replace("DUC-", "", 1)
        refs = sorted(uc_ref_map.get(alias, set()))
        if refs:
            uc["req_ids"] = refs
    return names


def _is_testable_requirement(text: str) -> tuple[bool, list]:
    t = str(text or "")
    measurable_patterns = [
        r"\b\d+(\.\d+)?\s*(ms|s|sec|seconds|minutes|min|hours|%)\b",
        r"\b(less than|greater than|at least|at most|within)\b\s*\d+",
        r"\b\d+\s*(users|requests|transactions|records|items)\b",
    ]
    vague_terms = [
        "fast", "user-friendly", "easy", "quick", "efficient", "robust", "scalable", "secure enough",
        "as soon as possible", "high performance", "good", "better",
    ]
    has_measure = any(re.search(p, t, flags=re.IGNORECASE) for p in measurable_patterns)
    found_vague = [v for v in vague_terms if re.search(rf"\b{re.escape(v)}\b", t, flags=re.IGNORECASE)]
    if has_measure and not found_vague:
        return True, []
    if has_measure and found_vague:
        return True, found_vague
    return False, found_vague


def _build_rtm_report(srs_data: dict, use_case_data: dict) -> dict:
    reqs = _extract_requirements_for_rtm(srs_data)
    textual = _extract_textual_usecases_for_rtm(use_case_data)
    diagram = _extract_diagram_usecases_for_rtm(use_case_data)
    rows = []
    covered_count = 0
    consistent_count = 0
    testable_count = 0
    used_textual_ids = set()
    used_diagram_ids = set()

    for req in reqs:
        req_norm = _normalize_req_id(req.get("req_id", ""))
        best_t = []
        best_d = []
        for uc in textual:
            score = _rtm_similarity(req["text"], uc["text"])
            uc_refs = {_normalize_req_id(x) for x in uc.get("req_ids", []) if _normalize_req_id(x)}
            explicit = req_norm and req_norm in uc_refs
            if score >= 0.12 or explicit:
                if explicit:
                    score = max(score, 0.95)
                best_t.append((score, uc))
        for du in diagram:
            score = _rtm_similarity(req["text"], du["text"])
            du_refs = {_normalize_req_id(x) for x in du.get("req_ids", []) if _normalize_req_id(x)}
            explicit = req_norm and req_norm in du_refs
            if score >= 0.08 or explicit:
                if explicit:
                    score = max(score, 0.95)
                best_d.append((score, du))
        best_t.sort(key=lambda x: x[0], reverse=True)
        best_d.sort(key=lambda x: x[0], reverse=True)
        best_t = best_t[:3]
        best_d = best_d[:3]

        textual_ids = [x[1]["id"] for x in best_t]
        textual_names = [x[1]["name"] for x in best_t]
        diagram_ids = [x[1]["id"] for x in best_d]
        diagram_names = [x[1]["name"] for x in best_d]
        for tid in textual_ids:
            used_textual_ids.add(tid)
        for did in diagram_ids:
            used_diagram_ids.add(did)

        covered = bool(textual_ids or diagram_ids)
        if covered:
            covered_count += 1
        consistency = "missing"
        if textual_ids and diagram_ids:
            textual_names_norm = {
                _canonical_uc_name(name) for name in textual_names if _canonical_uc_name(name)
            }
            diagram_names_norm = {
                _canonical_uc_name(name) for name in diagram_names if _canonical_uc_name(name)
            }
            name_overlap = bool(textual_names_norm.intersection(diagram_names_norm))
            textual_refs = {
                _normalize_req_id(x)
                for _, uc in best_t
                for x in uc.get("req_ids", [])
                if _normalize_req_id(x)
            }
            diagram_refs = {
                _normalize_req_id(x)
                for _, du in best_d
                for x in du.get("req_ids", [])
                if _normalize_req_id(x)
            }
            ref_overlap = bool(textual_refs.intersection(diagram_refs))
            consistency = "good" if (name_overlap or ref_overlap) else "partial"
        elif textual_ids or diagram_ids:
            consistency = "partial"
        if consistency == "good":
            consistent_count += 1
        testable, vague_terms = _is_testable_requirement(req["text"])
        if testable:
            testable_count += 1

        rows.append(
            {
                "req_id": req["req_id"],
                "requirement": req["text"],
                "type": req["type"],
                "textual_usecases": textual_ids,
                "textual_usecase_names": textual_names,
                "diagram_usecases": diagram_ids,
                "diagram_usecase_names": diagram_names,
                "coverage_status": "covered" if covered else "uncovered",
                "consistency_status": consistency,
                "testability_status": "testable" if testable else "needs_refinement",
                "notes": (
                    "Contains vague terms: " + ", ".join(vague_terms)
                    if vague_terms
                    else ("No linked use case found." if not covered else "Linked across available artifacts.")
                ),
            }
        )

    all_textual_ids = {u["id"] for u in textual}
    all_diagram_ids = {u["id"] for u in diagram}
    orphan_textual = sorted(list(all_textual_ids - used_textual_ids))
    orphan_diagram = sorted(list(all_diagram_ids - used_diagram_ids))
    total = len(reqs)
    return {
        "summary": {
            "total_requirements": total,
            "covered_requirements": covered_count,
            "coverage_ratio": (covered_count / total) if total else 0.0,
            "consistent_requirements": consistent_count,
            "consistency_ratio": (consistent_count / total) if total else 0.0,
            "testable_requirements": testable_count,
            "testability_ratio": (testable_count / total) if total else 0.0,
            "textual_usecases_total": len(textual),
            "diagram_usecases_total": len(diagram),
            "orphan_textual_usecases": orphan_textual,
            "orphan_diagram_usecases": orphan_diagram,
        },
        "rows": rows,
    }


# Initialize orchestrator
orchestrator = RequirementsOrchestrator()
# Ensure audio transcription is enabled for the API unless explicitly disabled elsewhere
try:
    orchestrator.processor.config.enable_whisper = True
except Exception:
    pass

# Helper: serialize SRS object to dict with hallucination analysis if present
def serialize_srs(srs):
    raw_text = getattr(srs, 'raw_text', None) or srs.sections.get('_raw_text')
    hallucination_analysis = srs.sections.get('_hallucination_analysis', {})
    return {
        'document_id': srs.document_id,
        'title': srs.title,
        'version': srs.version,
        'date': srs.date,
        'author': srs.author,
        'sections': srs.sections,
        'raw_text': raw_text,
        'hallucination_analysis': hallucination_analysis,
    }


def _extract_text_from_pdf(pdf_path: str) -> str:
    """Best-effort text extraction from PDF; falls back gracefully."""
    try:
        import PyPDF2
    except Exception:
        return ""


def _get_client_ip() -> str:
    xff = request.headers.get("X-Forwarded-For", "")
    if xff:
        return xff.split(",")[0].strip()
    return request.remote_addr or "unknown"


def _stable_text_fingerprint(text: str) -> str:
    try:
        import hashlib

        return hashlib.sha256(str(text or "").encode("utf-8", errors="ignore")).hexdigest()[:16]
    except Exception:
        return "na"


def _record_suspicious_input_event(reason: str, text: str = "") -> None:
    ip = _get_client_ip()
    now = time.time()
    with _SECURITY_LOCK:
        events = _SUSPICIOUS_EVENTS[ip]
        events.append(now)
        while events and (now - events[0]) > _SECURITY_WINDOW_SECONDS:
            events.popleft()
        if len(events) >= 6:
            _BLOCKED_UNTIL[ip] = now + _SECURITY_ATTACK_BLOCK_SECONDS
    logger.warning(
        "SECURITY suspicious_input ip=%s reason=%s fingerprint=%s",
        ip,
        reason,
        _stable_text_fingerprint(text),
    )


def _decode_obfuscated_variants(text: str) -> list[str]:
    """
    Generate alternate decoded forms to catch encoded prompt-injection payloads.
    """
    src = str(text or "")
    variants = [src]
    try:
        u1 = unquote_plus(src)
        if u1 != src:
            variants.append(u1)
            u2 = unquote_plus(u1)
            if u2 != u1:
                variants.append(u2)
    except Exception:
        pass

    # Decode long base64-like fragments inside text.
    for token in re.findall(r"[A-Za-z0-9+/=]{24,}", src):
        if len(token) % 4 != 0:
            continue
        try:
            raw = base64.b64decode(token, validate=True)
            decoded = raw.decode("utf-8", errors="ignore")
            if decoded and any(ch.isalpha() for ch in decoded):
                variants.append(decoded)
        except (binascii.Error, ValueError):
            continue
    return variants


def _normalize_for_security_scan(text: str) -> str:
    norm = str(text or "")
    norm = norm.replace("\x00", " ")
    norm = norm.replace("\u200b", "").replace("\u200c", "").replace("\u200d", "").replace("\ufeff", "")
    # Simple leetspeak normalization for common bypasses.
    norm = norm.translate(str.maketrans({"0": "o", "1": "i", "3": "e", "4": "a", "5": "s", "7": "t"}))
    norm = re.sub(r"[^a-zA-Z0-9\s]", " ", norm)
    norm = re.sub(r"\s+", " ", norm).strip().lower()
    return norm
    try:
        text_chunks = []
        with open(pdf_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                try:
                    text_chunks.append(page.extract_text() or "")
                except Exception:
                    continue
        return "\n".join(chunks for chunks in text_chunks if chunks)
    except Exception:
        return ""

# High-confidence prompt-injection / tokenizer abuse only.
# Kept intentionally narrow: phrases like "from now on", "system prompt" (UI copy), or
# "the system generates alerts" are valid requirements and must NOT match here.
SUSPICIOUS_PATTERNS = [
    r'\bignore\s+(?:all\s+)?previous\s+instructions?\b',
    r'\bignore\s+all\s+instructions?\b',
    r'\bdisregard\s+(?:all\s+)?(?:previous\s+)?instructions?\b',
    r'\bremove\s+(?:all\s+)?previous\s+instructions?\b',
    r'\bdelete\s+(?:all\s+)?previous\s+instructions?\b',
    r'\bdiscard\s+(?:all\s+)?previous\s+instructions?\b',
    r'\b(?:clear|erase)\s+all\s+previous\s+instructions?\b',
    r'\breplace\s+(?:all\s+)?previous\s+instructions?\b',
    r'\bdo\s+not\s+follow\s+(?:any\s+)?(?:previous\s+)?instructions?\b',
    r'\bforget\s+all\s+(?:previous\s+)?instructions?\b',
    r'\bforget\s+everything\b',
    r'\bprompt\s+injection\b',
    r'\bignore\s+the\s+prompt\b',
    r'\b(?:new|updated)\s+instructions\s*:\s*',
    r'\bact\s+as\s+if\b',
    r'\bpretend\s+to\s+be\b',
    r'\bpretend\s+you\s+are\b',
    r'\broleplay\s+as\b',
    r'<\|[^|]+\|>',
    r'\[INST\].*?\[/INST\]',
    r'<\|im_start\|>.*?<\|im_end\|>',
    r'<\|(user|assistant|system|eot_id)\|>',
]

def detect_prompt_injection(text: str) -> tuple[bool, list[str]]:
    """
    Detect potential prompt injection attempts in user input.
    
    Args:
        text: The text content to check for prompt injection patterns
    
    Returns:
        Tuple containing:
            - bool: True if suspicious patterns detected, False otherwise
            - list: List of detected suspicious patterns
    """
    if not text:
        return False, []

    detected_patterns: list[str] = []
    scan_variants = _decode_obfuscated_variants(text)
    normalized_variants = [_normalize_for_security_scan(v) for v in scan_variants]

    for variant in normalized_variants:
        for pattern in SUSPICIOUS_PATTERNS:
            if re.search(pattern, variant, re.IGNORECASE | re.MULTILINE):
                detected_patterns.append(pattern)

    # Heuristic for disguised imperative attacks ("ignore previous instructions" broken by symbols/spaces)
    joined = " ".join(normalized_variants)
    if (
        "ignore" in joined
        and "previous" in joined
        and ("instruction" in joined or "instructions" in joined)
    ):
        detected_patterns.append("heuristic:ignore_previous_instructions")

    # unique
    detected_patterns = sorted(set(detected_patterns))
    if detected_patterns:
        _record_suspicious_input_event("prompt_injection_detected", text)

    return bool(detected_patterns), detected_patterns

def sanitize_user_input(text: str, max_length: int = 10000) -> str:
    """
    Sanitize user input to prevent prompt injection attacks.
    
    Removes or neutralizes potentially dangerous content:
    - Markdown code blocks
    - Suspicious instruction patterns
    - Excessive special characters
    - Truncates to maximum length
    
    Args:
        text: Raw user input text
        max_length: Maximum allowed length (default: 10000 characters)
    
    Returns:
        Sanitized text safe for use in prompts
    """
    if not text:
        return ""
    
    # Remove markdown code blocks that could contain instructions
    text = re.sub(r'```[\s\S]*?```', '', text)
    text = re.sub(r'`[^`]+`', '', text)
    
    # Remove HTML-like tags that might be interpreted as instructions
    text = re.sub(r'<[^>]+>', '', text)
    
    # Remove special instruction markers
    text = re.sub(r'<\|.*?\|>', '', text)
    text = re.sub(r'\[INST\].*?\[/INST\]', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<\|im_start\|>.*?<\|im_end\|>', '', text, flags=re.DOTALL | re.IGNORECASE)
    
    # Remove excessive newlines (more than 2 consecutive)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Remove excessive whitespace
    text = re.sub(r' {3,}', ' ', text)
    
    # Truncate to maximum length
    if len(text) > max_length:
        text = text[:max_length].rsplit(' ', 1)[0]  # Cut at word boundary
    
    return text.strip()


class RequirementSecurityError(Exception):
    """Raised when structured requirement payloads fail injection checks (maps to HTTP 403)."""

    def __init__(self, payload: dict, status: int = 403):
        self.payload = payload
        self.status = status
        super().__init__(payload.get("error", "Rejected"))


# Keys whose string values are merged into LLM prompts or re-displayed — sanitize defensively.
_RESULT_TEXT_KEYS = frozenset(
    {
        "original_text",
        "content",
        "text",
        "description",
        "processing",
        "input",
        "output",
        "name",
        "title",
        "main_success_scenario",
    }
)


def _sanitize_requirement_dict(d: dict) -> dict:
    cleaned: dict = {}
    for k, v in d.items():
        if isinstance(v, str) and k in _RESULT_TEXT_KEYS:
            cleaned[k] = sanitize_user_input(v)
        elif k == "extracted_fields" and isinstance(v, dict):
            cleaned[k] = {
                ek: sanitize_user_input(ev) if isinstance(ev, str) else ev for ek, ev in v.items()
            }
        elif isinstance(v, dict):
            cleaned[k] = _sanitize_requirement_dict(v)
        elif isinstance(v, list):
            cleaned[k] = [
                _sanitize_requirement_dict(x) if isinstance(x, dict) else x for x in v
            ]
        else:
            cleaned[k] = v
    return cleaned


def sanitize_requirement_results(results: Any) -> list:
    """Strip hazardous markup from structured requirement payloads (direct API / tampered clients)."""
    if not isinstance(results, list):
        results = [results] if results is not None else []
    out: list = []
    for item in results:
        if isinstance(item, dict):
            out.append(_sanitize_requirement_dict(item))
        else:
            out.append(item)
    return out


def sanitize_project_info(project_info: Any) -> dict:
    """Sanitize string fields copied into generation prompts."""
    if not isinstance(project_info, dict):
        return {}
    pi = dict(project_info)
    for key in ("title", "author", "version"):
        if key in pi and isinstance(pi[key], str):
            pi[key] = sanitize_user_input(pi[key])[:4000]
    return pi


def _combined_requirement_text_from_results(results: list) -> str:
    parts: list[str] = []
    for item in results:
        if not isinstance(item, dict):
            continue
        t = str(
            item.get("original_text") or item.get("content") or item.get("text") or ""
        ).strip()
        if t:
            parts.append(t)
    return "\n".join(parts).strip()


def reject_payload_if_prompt_injection_in_results(results: list) -> Optional[dict]:
    """
    If combined requirement text matches injection heuristics, return an error body dict for JSON 403.
    """
    combined = _combined_requirement_text_from_results(results)
    if not combined:
        return None
    has_inj, _ = detect_prompt_injection(combined)
    if not has_inj:
        return None
    logger.warning(
        "SECURITY: blocked SRS payload — prompt-injection heuristics matched on combined requirement text"
    )
    return {
        "error": "Text content validation failed",
        "validation_errors": [
            "Input appears to contain prompt-injection patterns. Provide legitimate requirement content only."
        ],
        "security_issue": True,
    }


def validate_text_content(text: str) -> dict:
    """
    Validates text content against system requirements.
    
    Performs validation checks to ensure text content meets minimum quality standards:
    - Text must not be empty
    - Must contain at least one alphabetic character
    - Must have a minimum of 50 words
    - Must not be dominated by repeated tokens (anti-spam / anti-garbage)
    - Must not contain prompt-injection attempts
    
    Args:
        text: The text content to validate
    
    Returns:
        Dictionary containing:
            - 'valid' (bool): True if text passes all validation checks
            - 'errors' (list): List of error messages describing validation failures
    """
    errors = []
    
    if not text or not text.strip():
        return {'valid': False, 'errors': ['Text content is empty']}
    
    # Check for prompt injection attempts
    has_injection, detected_patterns = detect_prompt_injection(text)
    if has_injection:
        # Log security event with more context
        logger.warning(
            f"SECURITY ALERT: Prompt injection attempt detected. "
            f"Patterns: {detected_patterns}. "
            f"Input length: {len(text)} chars. "
            f"Fingerprint: {_stable_text_fingerprint(text)}"
        )
        errors.append(
            'Request rejected: disallowed instruction-hijack phrasing (e.g. ignore/remove/delete previous instructions, '
            'or model control tokens). Remove it and submit only genuine product requirements.'
        )
        return {
            'valid': False,
            'errors': errors,
            'security_issue': True
        }
    
    # Detect excessive repetition (e.g., "hi hi hi ..." spam)
    words = text.strip().split()
    if words:
        from collections import Counter
        counts = Counter(w.lower() for w in words if w.strip())
        total = sum(counts.values())
        most_common_word, freq = counts.most_common(1)[0]
        if total >= 20 and freq / total >= 0.6:
            errors.append(f"Input appears to be mostly repeated token '{most_common_word}' ({freq} of {total} tokens). Please provide meaningful requirements.")
            _record_suspicious_input_event("repeated_token_spam", text)
    
    # Require at least one alphabetic character (allow numbers/symbols but not only them)
    if not re.search(r'[A-Za-z]', text):
        errors.append('Text must include at least one alphabetic character (A-Z).')
    # Check minimum word count (50 words)
    words = text.strip().split()
    word_count = len([word for word in words if word])
    if word_count < 50:
        errors.append(f'Minimum 50 words required (current: {word_count} words)')
    
    return {
        'valid': len(errors) == 0,
        'errors': errors
    }


_SENSITIVE_OUTPUT_PATTERNS = [
    (re.compile(r"-----BEGIN [A-Z ]*PRIVATE KEY-----[\s\S]*?-----END [A-Z ]*PRIVATE KEY-----", re.IGNORECASE), "[REDACTED_PRIVATE_KEY]"),
    (re.compile(r"\bsk-[A-Za-z0-9]{20,}\b"), "[REDACTED_API_KEY]"),
    (re.compile(r"\bAKIA[0-9A-Z]{16}\b"), "[REDACTED_AWS_KEY]"),
    (re.compile(r"\b(?:api[_-]?key|access[_-]?token|secret|password)\s*[:=]\s*[^\s,;]{6,}", re.IGNORECASE), "[REDACTED_SECRET]"),
    (re.compile(r"(?i)\b(system prompt|developer instructions?|hidden instructions?)\b"), "[REDACTED_INTERNAL_INSTRUCTIONS]"),
]


def _sanitize_output_text(text: str) -> str:
    out = str(text or "")
    for pattern, replacement in _SENSITIVE_OUTPUT_PATTERNS:
        out = pattern.sub(replacement, out)
    out = re.sub(r"\n{4,}", "\n\n\n", out)
    return out


def _sanitize_output_payload(obj: Any) -> Any:
    if isinstance(obj, str):
        return _sanitize_output_text(obj)
    if isinstance(obj, list):
        return [_sanitize_output_payload(x) for x in obj]
    if isinstance(obj, dict):
        return {k: _sanitize_output_payload(v) for k, v in obj.items()}
    return obj


def build_clarification_payload(text: str) -> dict:
    """
    Build clarification suggestions for vague requirements.
    Uses ambiguity detection + refinement modules and computes a soft quality score.
    """
    detector = AmbiguityDetector()
    refiner = RequirementRefiner()
    analysis = detector.analyze_requirement(text)
    refined = refiner.refine(text)
    details = analysis.get("details", [])

    ambiguous_count = len(details)
    weak_modals = sum(1 for item in details if item.get("category") == "weak_modal")
    vague_quantifiers = sum(1 for item in details if item.get("category") == "vague_quantifier")

    # Missing details prompts inferred from text coverage.
    lowered = text.lower()
    add_suggestions = []
    if not re.search(r"\b(user|admin|customer|operator|system)\b", lowered):
        add_suggestions.append("Specify who performs each action (e.g., User, Admin, System).")
    if not re.search(r"\b(within|under|less than|at least|%|seconds|ms|hours|days)\b", lowered):
        add_suggestions.append("Add measurable constraints (e.g., response time, throughput, uptime).")
    if not re.search(r"\b(if|when|on failure|error|invalid|timeout)\b", lowered):
        add_suggestions.append("Add exception/error-handling behavior for failure scenarios.")
    if not re.search(
        r"\b(authentication|authorization|login|password|encrypt|tls|ssl|audit|gdpr|privacy|pii|rbac)\b",
        lowered,
    ):
        add_suggestions.append(
            "Clarify security: authentication/authorization model, sensitive data handling, and audit expectations."
        )
    if not re.search(r"\b(availability|uptime|backup|recovery|disaster|rto|rpo|redundan)\b", lowered):
        add_suggestions.append(
            "Specify reliability targets: expected uptime, backup frequency, and recovery expectations where relevant."
        )
    if not re.search(r"\b(api|interface|integrat|webhook|rest|export|import|third[- ]party)\b", lowered):
        add_suggestions.append(
            "Name external systems or interfaces (APIs, files, services) the product must connect to—core functional scope."
        )
    if not re.search(r"\b(role|permission|admin|customer|actor|stakeholder)\b", lowered):
        add_suggestions.append(
            "List primary actors/roles and what each may do (helps separate FRs per user class)."
        )

    remove_suggestions = []
    for item in details:
        remove_suggestions.append(
            f"Replace '{item.get('word')}' with measurable wording: {item.get('replacement')}"
        )

    # Soft quality score: lower penalty than strict gating.
    score = 1.0
    score -= min(0.45, ambiguous_count * 0.06)
    score -= min(0.20, weak_modals * 0.03)
    score -= min(0.20, vague_quantifiers * 0.04)
    score -= min(0.15, len(add_suggestions) * 0.05)
    clarification_score = max(0.0, min(1.0, round(score, 3)))

    unresolved_count = ambiguous_count + len(add_suggestions)
    warning_level = "low"
    if unresolved_count >= 6:
        warning_level = "high"
    elif unresolved_count >= 3:
        warning_level = "medium"

    return {
        "input_text": text,
        "highlighted_text": analysis.get("highlighted_text", text),
        "suggested_rewrite": analysis.get("suggestion", text),
        "ambiguities": details,
        "clarification_questions": refined.get("clarification_questions", []),
        "add_suggestions": add_suggestions,
        "remove_suggestions": remove_suggestions[:20],
        "clarification_score": clarification_score,
        "unresolved_items": unresolved_count,
        "warning_level": warning_level,
        "structured_requirements": refined.get("structured_requirements", {}),
    }


def _clarification_fallback_payload(text: str, reason: str = "") -> dict:
    """
    Minimal deterministic clarification payload used when the full analyzer fails.
    Keeps frontend flow working instead of showing a hard failure banner.
    """
    lowered = (text or "").lower()
    add_suggestions = []
    if not re.search(r"\b(user|admin|customer|operator|system|doctor|patient)\b", lowered):
        add_suggestions.append("Specify the primary actor for each requirement (e.g., User, Admin, System).")
    if not re.search(r"\b(within|under|less than|at least|seconds|ms|minutes|hours|days|%)\b", lowered):
        add_suggestions.append("Add measurable targets (response time, throughput, or uptime) where applicable.")
    if not re.search(r"\b(if|when|error|invalid|timeout|failure)\b", lowered):
        add_suggestions.append("Add failure/exception behavior (what happens when operations fail).")
    if not re.search(r"\b(security|authentication|authorization|encrypt|privacy|audit)\b", lowered):
        add_suggestions.append("Clarify security/privacy expectations for access and sensitive data.")

    # conservative score so UI warns gently
    unresolved = len(add_suggestions)
    warning = "high" if unresolved >= 5 else "medium" if unresolved >= 3 else "low"
    score = max(0.35, 0.85 - unresolved * 0.08)
    payload = {
        "input_text": text,
        "highlighted_text": text,
        "suggested_rewrite": text,
        "ambiguities": [],
        "clarification_questions": [],
        "add_suggestions": add_suggestions,
        "remove_suggestions": [],
        "clarification_score": round(score, 3),
        "unresolved_items": unresolved,
        "warning_level": warning,
        "structured_requirements": {},
        "fallback_used": True,
    }
    if reason:
        payload["fallback_reason"] = reason
    return payload


def _copilot_aspect_title(category: str) -> str:
    """Human-readable aspect for suggestion cards (FR/NFR-style angles)."""
    labels = {
        "weak_modal": "Obligation & testability (modal verbs)",
        "vague_quantifier": "Time & commitment",
        "performance": "Performance (NFR)",
        "security": "Security & privacy (NFR)",
        "usability": "Usability & UX (NFR)",
        "reliability": "Reliability & recovery (NFR)",
        "scalability": "Scalability (NFR)",
        "time": "Timeliness & deadlines",
        "size": "Resource constraints",
    }
    return labels.get(category, "Clarity & precision")


def _copilot_fallback_from_clarification(clarification: dict) -> dict:
    """
    Deterministic copilot: avoid duplicate cards (e.g. three identical 'should' rows).
    Mix ambiguity categories with add_suggestions (scope, NFR, security) and structured hints.
    """
    rewrite = clarification.get("suggested_rewrite", "")
    ambiguities = list(clarification.get("ambiguities") or [])

    # One ambiguity per category first (breadth), then fill by position if needed.
    by_category: dict = {}
    for item in ambiguities:
        cat = item.get("category") or "general"
        if cat not in by_category:
            by_category[cat] = item

    priority_cats = [
        "weak_modal",
        "vague_quantifier",
        "security",
        "performance",
        "usability",
        "reliability",
        "scalability",
        "time",
        "size",
    ]
    ordered_amb: list = []
    for cat in priority_cats:
        if cat in by_category:
            ordered_amb.append(by_category[cat])
    for cat, item in by_category.items():
        if item not in ordered_amb:
            ordered_amb.append(item)

    suggestion_cards: list = []
    seen_keys: set = set()

    for item in ordered_amb:
        if len(suggestion_cards) >= 3:
            break
        word = item.get("word", "").strip() or "term"
        rep = item.get("replacement", "measurable criteria")
        cat = item.get("category", "general")
        aspect = _copilot_aspect_title(cat)
        key = f"{cat}:{word.lower()}"
        if key in seen_keys:
            continue
        seen_keys.add(key)
        suggestion_cards.append(
            {
                "title": f"{aspect}: refine '{word}'",
                "rewrite": rewrite,
                "reason": f"[{cat}] '{word}' is imprecise. Prefer testable wording such as: {rep}.",
            }
        )

    add_list = list(clarification.get("add_suggestions") or [])
    add_labels = [
        ("Actors & scope (FR)", r"who performs|actor|role|User|Admin|System"),
        ("Measurable targets (NFR)", r"measurable|response time|throughput|uptime"),
        ("Failure handling", r"failure|error|exception|timeout"),
        ("Security & privacy (NFR)", r"security|authentication|encrypt|audit|privacy"),
        ("Reliability & ops (NFR)", r"availability|backup|recovery|uptime"),
        ("Interfaces & integrations (FR)", r"interface|API|integrat"),
    ]

    for title_hint, pattern in add_labels:
        if len(suggestion_cards) >= 3:
            break
        for sug in add_list:
            if not isinstance(sug, str) or not sug.strip():
                continue
            if re.search(pattern, sug, re.IGNORECASE):
                fk = f"add:{title_hint}"
                if fk in seen_keys:
                    continue
                seen_keys.add(fk)
                suggestion_cards.append(
                    {
                        "title": title_hint,
                        "rewrite": rewrite,
                        "reason": sug.strip(),
                    }
                )
                break

    for sug in add_list:
        if len(suggestion_cards) >= 3:
            break
        if not isinstance(sug, str) or not sug.strip():
            continue
        fk = f"addtext:{sug[:48]}"
        if fk in seen_keys:
            continue
        seen_keys.add(fk)
        suggestion_cards.append(
            {
                "title": "Gap to address (requirements engineering)",
                "rewrite": rewrite,
                "reason": sug.strip(),
            }
        )

    structured = clarification.get("structured_requirements") or {}
    frs = structured.get("functional_requirements") or []
    nfrs = structured.get("non_functional_requirements") or []

    if len(suggestion_cards) < 3 and frs:
        st = frs[0] if isinstance(frs[0], dict) else {}
        hint = (st.get("refined_text") or st.get("source_text") or "")[:280]
        if hint:
            suggestion_cards.append(
                {
                    "title": "Functional requirement angle (draft)",
                    "rewrite": rewrite,
                    "reason": f"From your text, a clearer FR-style phrasing could be: {hint}",
                }
            )

    if len(suggestion_cards) < 3 and nfrs:
        st = nfrs[0] if isinstance(nfrs[0], dict) else {}
        hint = (st.get("refined_text") or st.get("source_text") or "")[:280]
        if hint:
            suggestion_cards.append(
                {
                    "title": "Non-functional requirement angle (draft)",
                    "rewrite": rewrite,
                    "reason": f"Strengthen measurability for quality attributes: {hint}",
                }
            )

    while len(suggestion_cards) < 3:
        suggestion_cards.append(
            {
                "title": "Add measurable acceptance criteria",
                "rewrite": rewrite,
                "reason": "Specific metrics (time, throughput, accuracy, uptime) improve testability and reduce ambiguity.",
            }
        )

    followups = clarification.get("clarification_questions") or []
    # Prefer a non-weak_modal question when possible for variety.
    q_alt = next((q for q in followups if "weak_modal" not in str(q)), None)
    primary_q = q_alt or (followups[0] if followups else None)

    return {
        "question": primary_q
        or "What measurable targets (performance, security, availability) matter most for this system?",
        "suggestions": suggestion_cards[:3],
        "source": "rule_fallback",
    }


def _copilot_llm_response(user_text: str, clarification: dict) -> dict | None:
    """
    Optional LLM enhancer for copilot suggestions.
    Returns None on failure so caller can fallback safely.
    """
    token = os.environ.get("REPLICATE_API_TOKEN", "").strip()
    model_name = os.environ.get(
        "REPLICATE_MODEL",
        "ar18now/qwen2:e2488e00bc2be9f83f548b6f1591c4dcc69cd6dc5e7a82ceb4968dc209ebd420",
    )
    if not token:
        return None

    prompt = f"""You are a requirements clarification copilot.
Given user requirements text and rule-based ambiguity analysis, return STRICT JSON only:
{{
  "question": "one targeted follow-up question",
  "suggestions": [
    {{"title":"...", "rewrite":"...", "reason":"..."}},
    {{"title":"...", "rewrite":"...", "reason":"..."}},
    {{"title":"...", "rewrite":"...", "reason":"..."}}
  ]
}}
Rules:
- Suggestions must cover DIFFERENT angles when possible: e.g. (1) functional/scope/actors, (2) non-functional measurable targets, (3) security/privacy OR interfaces OR error handling.
- Do NOT output three cards about the same word (e.g. three times "should").
- Keep suggestions grounded in user text only.
- Do not invent new product features.
- Prefer measurable wording.
- Max 3 suggestions.

USER_TEXT:
{user_text}

RULE_ANALYSIS:
{json.dumps(clarification, ensure_ascii=False)[:6000]}
"""
    try:
        os.environ["REPLICATE_API_TOKEN"] = token
        output = replicate.run(
            model_name,
            input={
                "prompt": prompt,
                "max_new_tokens": 450,
                "temperature": 0.2,
                "top_p": 0.9,
                "repetition_penalty": 1.05,
            },
            timeout=60,
        )
        text = "".join(output) if isinstance(output, list) else str(output)
        match = re.search(r"\{[\s\S]*\}", text)
        if not match:
            return None
        parsed = json.loads(match.group(0))
        if not isinstance(parsed, dict):
            return None
        suggestions = parsed.get("suggestions") or []
        cleaned = []
        for s in suggestions[:3]:
            if not isinstance(s, dict):
                continue
            cleaned.append(
                {
                    "title": str(s.get("title", "Suggestion")).strip()[:120],
                    "rewrite": str(s.get("rewrite", "")).strip()[:3000],
                    "reason": str(s.get("reason", "")).strip()[:600],
                }
            )
        if not cleaned:
            return None
        return {
            "question": str(parsed.get("question", "")).strip()[:300],
            "suggestions": cleaned,
            "source": "llm",
        }
    except Exception:
        return None


def _build_next_question_queue(clarification: dict) -> list[str]:
    base_questions = clarification.get("clarification_questions") or []
    if not base_questions:
        base_questions = [
            "Can you add measurable performance targets (e.g., response time/throughput)?",
            "Can you clarify security/privacy constraints for this system?",
            "Which actor performs each major action (User/Admin/System)?",
        ]
    # De-duplicate similar weak_modal questions (same template).
    out: list = []
    seen: set = set()
    for q in base_questions:
        q = str(q).strip()
        if not q or q in seen:
            continue
        seen.add(q)
        out.append(q)
    extras = [
        "What interfaces or external systems must this product integrate with?",
        "What are the expected availability and recovery expectations (uptime, backup)?",
    ]
    for e in extras:
        if len(out) >= 8:
            break
        if e not in seen:
            seen.add(e)
            out.append(e)
    return out[:8]

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'service': 'Requirements Engineering API'
    })


@app.route('/api/clarify-requirements', methods=['POST'])
def clarify_requirements():
    """Analyze requirements and return clarification suggestions before SRS generation."""
    try:
        data = request.get_json() or {}
        content = data.get('content', '')
        if not content or not str(content).strip():
            return jsonify({'error': 'No content provided'}), 400

        sanitized_content = sanitize_user_input(str(content))
        if not sanitized_content:
            return jsonify({'error': 'Content is invalid after sanitization'}), 400

        has_inj, _ = detect_prompt_injection(sanitized_content)
        if has_inj:
            logger.warning("clarify-requirements: blocked prompt-injection heuristics")
            return jsonify(
                {
                    'error': 'Text content validation failed',
                    'validation_errors': [
                        'Input appears to contain prompt-injection patterns. Please provide valid requirements instead.'
                    ],
                    'security_issue': True,
                }
            ), 403

        try:
            payload = build_clarification_payload(sanitized_content)
        except Exception as inner_err:
            logger.warning(f"clarify-requirements fallback used: {inner_err}")
            payload = _clarification_fallback_payload(
                sanitized_content,
                reason=f"Analyzer fallback: {inner_err}",
            )
        return jsonify(payload)
    except Exception as e:
        logger.error(f"Error clarifying requirements: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/clarification-copilot', methods=['POST'])
def clarification_copilot():
    """
    Hybrid copilot:
    - deterministic ambiguity/refinement baseline
    - optional LLM enhancement
    - safe fallback if LLM unavailable/fails
    """
    try:
        data = request.get_json() or {}
        content = str(data.get("content", "")).strip()
        if not content:
            return jsonify({"error": "No content provided"}), 400

        sanitized = sanitize_user_input(content)
        if not sanitized:
            return jsonify({"error": "Content is invalid after sanitization"}), 400

        has_inj, _ = detect_prompt_injection(sanitized)
        if has_inj:
            logger.warning("clarification-copilot: blocked prompt-injection heuristics")
            return jsonify(
                {
                    "error": "Text content validation failed",
                    "validation_errors": [
                        "Input appears to contain prompt-injection patterns. Please provide valid requirements instead."
                    ],
                    "security_issue": True,
                }
            ), 403

        clarification = build_clarification_payload(sanitized)
        copilot = _copilot_llm_response(sanitized, clarification) or _copilot_fallback_from_clarification(clarification)

        return jsonify(
            {
                "copilot": copilot,
                "clarification_score": clarification.get("clarification_score", 0),
                "unresolved_items": clarification.get("unresolved_items", 0),
                "suggested_rewrite": clarification.get("suggested_rewrite", sanitized),
                "question_queue": _build_next_question_queue(clarification),
            }
        )
    except Exception as e:
        logger.error(f"Error in clarification copilot: {str(e)}")
        return jsonify({"error": str(e)}), 500


@app.route('/api/clarification-copilot-turn', methods=['POST'])
def clarification_copilot_turn():
    """
    Conversational turn endpoint:
    - Takes latest answer and remaining question queue
    - Returns next question + refreshed suggestions
    """
    try:
        data = request.get_json() or {}
        content = str(data.get("content", "")).strip()
        latest_answer = str(data.get("latest_answer", "")).strip()
        question_queue = data.get("question_queue") or []
        if not isinstance(question_queue, list):
            question_queue = []
        question_queue = [str(q).strip() for q in question_queue if str(q).strip()]
        if content:
            merged_text = f"{content}\n\nClarification Answer:\n{latest_answer}" if latest_answer else content
        else:
            merged_text = latest_answer
        if not merged_text.strip():
            return jsonify({"error": "No content provided"}), 400

        sanitized = sanitize_user_input(merged_text)
        if not sanitized:
            return jsonify({"error": "Content is invalid after sanitization"}), 400

        has_inj, _ = detect_prompt_injection(sanitized)
        if has_inj:
            logger.warning("clarification-copilot-turn: blocked prompt-injection heuristics")
            return jsonify(
                {
                    "error": "Text content validation failed",
                    "validation_errors": [
                        "Input appears to contain prompt-injection patterns. Please provide valid requirements instead."
                    ],
                    "security_issue": True,
                }
            ), 403

        clarification = build_clarification_payload(sanitized)
        copilot = _copilot_llm_response(sanitized, clarification) or _copilot_fallback_from_clarification(clarification)

        # consume first queued question and return next
        if question_queue:
            question_queue = question_queue[1:]
        next_question = question_queue[0] if question_queue else copilot.get("question", "")

        return jsonify(
            {
                "copilot": {
                    "question": next_question,
                    "suggestions": copilot.get("suggestions", []),
                    "source": copilot.get("source", "rule_fallback"),
                },
                "question_queue": question_queue,
                "clarification_score": clarification.get("clarification_score", 0),
                "unresolved_items": clarification.get("unresolved_items", 0),
                "suggested_rewrite": clarification.get("suggested_rewrite", sanitized),
            }
        )
    except Exception as e:
        logger.error(f"Error in clarification copilot turn: {str(e)}")
        return jsonify({"error": str(e)}), 500


def _coerce_results_list_for_srs(results_data: object) -> list:
    """Match frontend buildRequirementsArray — shape passed to SRSModelGenerator."""
    if results_data is None:
        return []
    if isinstance(results_data, list):
        return results_data
    if isinstance(results_data, dict):
        if isinstance(results_data.get("results"), list):
            return results_data["results"]
        if results_data.get("status"):
            return [results_data]
        return [results_data]
    return [results_data]


def _validate_generation_payload(results: list, project_info: dict) -> Optional[dict]:
    if not isinstance(results, list) or not results:
        return {"error": "Invalid generation payload: results must be a non-empty list."}
    if len(results) > 128:
        return {"error": "Invalid generation payload: too many requirement items."}
    for idx, item in enumerate(results):
        if not isinstance(item, dict):
            return {"error": f"Invalid generation payload: item {idx} must be an object."}
        txt = str(item.get("original_text") or item.get("content") or item.get("text") or "")
        if len(txt) > 20000:
            return {"error": f"Invalid generation payload: item {idx} is too large."}
    if project_info is not None and not isinstance(project_info, dict):
        return {"error": "Invalid generation payload: project_info must be an object."}
    return None


def _generate_srs_document(results: list, project_info: dict, mode_override=None):
    """
    Runtime generation switch.
    Env variables:
      - SRS_GENERATION_MODE=rag|model      (default: model)
      - RAG_KB_PATH=<path>                 (optional; default auto-discovery)
      - RAG_TOP_K=<int>                    (default: 6)
      - RAG_VECTOR_BACKEND=faiss|chroma    (default: faiss)
    """
    mode = str(mode_override or os.environ.get("SRS_GENERATION_MODE", "model")).strip().lower()
    if mode not in {"rag", "model"}:
        logger.warning("Unknown SRS_GENERATION_MODE=%s; using model.", mode)
        mode = "model"

    if mode == "rag":
        kb_path = str(os.environ.get("RAG_KB_PATH", "")).strip()
        vector_backend = str(os.environ.get("RAG_VECTOR_BACKEND", "faiss")).strip().lower() or "faiss"
        try:
            top_k = int(os.environ.get("RAG_TOP_K", "6"))
        except Exception:
            top_k = 6
        top_k = max(1, min(20, top_k))

        try:
            rag = RAGSRSGenerator(vector_backend=vector_backend)
            if kb_path:
                loaded_docs = rag.load_knowledge_base(kb_path)
                logger.info(
                    "SRS generation mode=rag backend=%s top_k=%s kb_path=%s loaded_docs=%s",
                    vector_backend,
                    top_k,
                    kb_path,
                    loaded_docs,
                )
            else:
                loaded_docs = rag.load_default_knowledge_base(str(Path(__file__).resolve().parent))
                logger.info(
                    "SRS generation mode=rag backend=%s top_k=%s kb_path=<default> loaded_docs=%s",
                    vector_backend,
                    top_k,
                    loaded_docs,
                )

            if loaded_docs <= 0:
                logger.warning("RAG mode requested but KB is empty; falling back to model mode.")
                model_gen = SRSModelGenerator()
                return model_gen.generate_srs(results, project_info)

            return rag.generate(results, project_info=project_info, top_k=top_k)
        except Exception as e:
            logger.error("RAG generation failed (%s). Falling back to model mode.", e, exc_info=True)
            model_gen = SRSModelGenerator()
            return model_gen.generate_srs(results, project_info)

    logger.info("SRS generation mode=model")
    model_gen = SRSModelGenerator()
    return model_gen.generate_srs(results, project_info)


def _build_srs_dict_from_results(results: list, project_info: dict, mode_override=None) -> dict:
    """Build SRS JSON payload (same contract as POST /api/generate-srs)."""
    if not results:
        raise ValueError("No results provided")
    if not isinstance(results, list):
        results = [results]

    results = sanitize_requirement_results(results)
    project_info = sanitize_project_info(project_info)
    payload_err = _validate_generation_payload(results, project_info)
    if payload_err:
        raise ValueError(payload_err["error"])
    inj = reject_payload_if_prompt_injection_in_results(results)
    if inj:
        raise RequirementSecurityError(inj, 403)

    logger.info(f"Generating SRS with {len(results)} result(s)")
    logger.debug(f"Results sample: {results[0] if results else 'No results'}")
    logger.debug(f"Project info: {project_info}")

    resolved_mode = str(mode_override or os.environ.get("SRS_GENERATION_MODE", "model")).strip().lower()
    if resolved_mode not in {"rag", "model"}:
        resolved_mode = "model"

    srs = _generate_srs_document(results, project_info, mode_override=resolved_mode)

    logger.info(f"SRS generated successfully: {srs.document_id}")

    return _srs_document_to_api_dict(srs, results, project_info, resolved_mode)


def _sse_srs_event(payload: dict) -> str:
    """One Server-Sent Event frame (JSON payload)."""
    return f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"


def _srs_document_to_api_dict(srs, results: list, project_info: dict, resolved_mode: str) -> dict:
    """Build the same JSON contract as POST /api/generate-srs from an in-memory SRS document."""
    raw_text = getattr(srs, "raw_text", None) or srs.sections.get("_raw_text")
    hallucination_analysis = srs.sections.get("_hallucination_analysis", {})

    source_text = "\n".join(
        str(item.get("original_text") or item.get("content") or item.get("text") or "")
        for item in results
        if isinstance(item, dict)
    ).strip()
    manual_metrics = ManualMetricsEngine().evaluate(
        source_requirements_text=source_text,
        structured_requirements={},
        srs_sections=srs.sections,
        srs_text=raw_text or json.dumps(srs.sections, ensure_ascii=False),
    )
    conflict_analysis = ConflictMetric().analyze(source_text)
    nfr_analysis = NFRSpecificityMetric().analyze(source_text, srs.sections)
    style_analysis = ProfessionalStyleMetric().analyze(
        raw_text or json.dumps(srs.sections, ensure_ascii=False), srs.sections
    )
    srs_dict = {
        "document_id": srs.document_id,
        "title": srs.title,
        "version": srs.version,
        "date": srs.date,
        "author": srs.author,
        "sections": srs.sections,
        "raw_text": raw_text,
        "hallucination_analysis": hallucination_analysis,
        "verification_report": {
            "manual_metrics": manual_metrics,
            "conflict_analysis": conflict_analysis,
            "nfr_specificity_analysis": nfr_analysis,
            "professional_style_analysis": style_analysis,
            "model_limitations": [
                "LLM output can still contain unsupported assumptions; review required.",
                "Ambiguity in input can propagate into vague FR/NFR unless clarified.",
                "Domain-specific NFR quality depends on source requirement specificity.",
                "Prompt-injection sanitization is heuristic and not a formal guarantee.",
            ],
        },
        "generation_meta": {
            "mode": resolved_mode,
        },
    }

    if hallucination_analysis.get("has_hallucinations"):
        logger.warning(
            f"SRS {srs.document_id} has potential hallucinations. "
            f"Confidence: {hallucination_analysis.get('confidence_score', 'N/A')}"
        )

    return _sanitize_output_payload(srs_dict)


@app.route('/api/process-single', methods=['POST'])
def process_single_requirement():
    """Process a single requirement from text input"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        input_type = data.get('type', 'text')
        content = data.get('content', '')
        project_info = data.get('project_info', {})
        
        if not content:
            return jsonify({'error': 'No content provided'}), 400
        
        # Sanitize user input to prevent prompt injection
        original_length = len(content)
        sanitized_content = sanitize_user_input(content)
        sanitized_length = len(sanitized_content)
        
        # Log if significant content was removed during sanitization
        if original_length > sanitized_length + 100:  # More than 100 chars removed
            logger.info(
                f"Input sanitization removed {original_length - sanitized_length} characters. "
                f"Original: {original_length} chars, Sanitized: {sanitized_length} chars"
            )
        
        if not sanitized_content or len(sanitized_content.strip()) < 10:
            logger.warning(f"Content rejected after sanitization. Original length: {original_length}, Sanitized length: {sanitized_length}")
            return jsonify({'error': 'Content is invalid or too short after sanitization'}), 400
        
        # Validate content (includes prompt injection detection)
        validation = validate_text_content(sanitized_content)
        if not validation['valid']:
            error_status = 403 if validation.get('security_issue') else 400
            return jsonify({
                'error': 'Text content validation failed', 
                'validation_errors': validation['errors']
            }), error_status

        # Prepare input data with sanitized content
        if input_type == 'text':
            input_data = {'type': 'text', 'content': sanitized_content}
        else:
            return jsonify({'error': 'Unsupported input type'}), 400
        
        # Process the requirement
        result = orchestrator.process_single_requirement(input_data)
        
        # Add project info to result
        result['project_info'] = project_info
        
        return jsonify(result)
        
    except Exception as e:
        logger.error(f"Error processing single requirement: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/process-and-generate-srs', methods=['POST'])
def process_and_generate_srs():
    """
    Single HTTP request: process text requirements then generate SRS on the server.
    Avoids two separate /process-single + /generate-srs round-trips (and duplicate billed calls).
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400

        input_type = data.get('type', 'text')
        content = data.get('content', '')
        project_info = data.get('project_info', {})

        if input_type != 'text':
            return jsonify({'error': 'Only type=text is supported for this endpoint'}), 400
        if not content:
            return jsonify({'error': 'No content provided'}), 400

        original_length = len(content)
        sanitized_content = sanitize_user_input(content)
        sanitized_length = len(sanitized_content)

        if original_length > sanitized_length + 100:
            logger.info(
                f"Input sanitization removed {original_length - sanitized_length} characters. "
                f"Original: {original_length} chars, Sanitized: {sanitized_length} chars"
            )

        if not sanitized_content or len(sanitized_content.strip()) < 10:
            logger.warning(
                f"Content rejected after sanitization. Original length: {original_length}, "
                f"Sanitized length: {sanitized_length}"
            )
            return jsonify({'error': 'Content is invalid or too short after sanitization'}), 400

        validation = validate_text_content(sanitized_content)
        if not validation['valid']:
            error_status = 403 if validation.get('security_issue') else 400
            return jsonify({
                'error': 'Text content validation failed',
                'validation_errors': validation['errors']
            }), error_status

        input_data = {'type': 'text', 'content': sanitized_content}
        result = orchestrator.process_single_requirement(input_data)
        result['project_info'] = project_info

        results_list = _coerce_results_list_for_srs(result)
        if not results_list:
            return jsonify({'error': 'Processing produced no results', 'processing': result}), 500

        try:
            srs_dict = _build_srs_dict_from_results(results_list, project_info)
        except RequirementSecurityError as sec:
            return jsonify(sec.payload), sec.status
        except Exception as srs_e:
            logger.error(f"SRS generation failed after processing: {srs_e}")
            return jsonify({
                'processing': result,
                'srs': None,
                'srs_error': str(srs_e),
            }), 200

        return jsonify({
            'processing': result,
            'srs': srs_dict,
        })

    except Exception as e:
        logger.error(f"Error in process-and-generate-srs: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/process-audio', methods=['POST'])
def process_audio_requirement():
    """Process a single requirement from audio recording"""
    temp_file_path = None
    try:
        audio_file = request.files.get('audio')
        project_info_str = request.form.get('project_info', '{}')
        
        if not audio_file:
            logger.error("No audio file provided in request")
            return jsonify({'error': 'No audio file provided'}), 400
        
        logger.info(f"Received audio file: {audio_file.filename}, Content-Type: {audio_file.content_type}, Size: {audio_file.content_length} bytes")
        
        # Parse project info
        try:
            project_info = json.loads(project_info_str) if project_info_str else {}
        except json.JSONDecodeError as e:
            logger.warning(f"Failed to parse project_info JSON: {e}, using empty dict")
            project_info = {}
        
        # Determine file extension from content type or filename
        original_filename = audio_file.filename or 'recording'
        file_ext = None
        
        # Try to get extension from filename
        if '.' in original_filename:
            file_ext = original_filename.rsplit('.', 1)[1].lower()
        
        # If no extension or unsupported, try to determine from content type
        if not file_ext or file_ext not in ['wav', 'webm', 'mp4', 'm4a', 'ogg', 'mp3', 'flac']:
            content_type = audio_file.content_type or ''
            if 'webm' in content_type or 'webm' in original_filename.lower():
                file_ext = 'webm'
            elif 'mp4' in content_type or 'm4a' in content_type or 'mp4' in original_filename.lower():
                file_ext = 'm4a'
            elif 'ogg' in content_type or 'ogg' in original_filename.lower():
                file_ext = 'ogg'
            elif 'wav' in content_type or 'wav' in original_filename.lower():
                file_ext = 'wav'
            else:
                # Default to webm (most common for browser recordings)
                file_ext = 'webm'
                logger.info(f"Unknown audio format, defaulting to webm")
        
        # Create temporary file with appropriate extension
        suffix = f'.{file_ext}'
        logger.info(f"Creating temp file with extension: {suffix}")
        
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temp_file:
            try:
                audio_file.save(temp_file.name)
                temp_file_path = temp_file.name
                logger.info(f"Audio file saved to: {temp_file_path}, Size: {os.path.getsize(temp_file_path)} bytes")
            except Exception as save_error:
                logger.error(f"Error saving audio file: {save_error}")
                return jsonify({'error': f'Failed to save audio file: {str(save_error)}'}), 500
        
        # Verify file was saved and has content
        if not os.path.exists(temp_file_path):
            logger.error(f"Temp file was not created: {temp_file_path}")
            return jsonify({'error': 'Failed to save audio file'}), 500
        
        file_size = os.path.getsize(temp_file_path)
        if file_size == 0:
            logger.error(f"Audio file is empty: {temp_file_path}")
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
            return jsonify({'error': 'Audio file is empty. Please record again.'}), 400

        # Reject very short audio (< 0.5s)
        try:
            duration = orchestrator.processor._get_audio_duration(temp_file_path)
            if duration <= 0.5:
                logger.error(f"Audio file too short ({duration}s): {temp_file_path}")
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                return jsonify({'error': 'Audio is too short. Please record at least 1 second.'}), 400
        except Exception as dur_err:
            logger.warning(f"Failed to measure duration for {temp_file_path}: {dur_err}")
        
        logger.info(f"Processing audio file: {temp_file_path} ({file_size} bytes)")
        
        # Process the audio requirement
        input_data = {'type': 'audio', 'file_path': temp_file_path}
        result = orchestrator.process_single_requirement(input_data)

        # If processing failed (e.g., whisper disabled or other error), return 400
        if result.get('status') == 'failed':
            error_msg = result.get('error', 'Audio processing failed')
            logger.error(f"Audio processing failed: {error_msg}")
            return jsonify({'error': error_msg}), 400
        
        # Validate the transcribed text
        if result.get('status') == 'completed':
            transcribed_text = result.get('original_text', '')
            logger.info(f"Transcription completed, text length: {len(transcribed_text)} characters")
        
        if not transcribed_text or not transcribed_text.strip():
            logger.error("Transcription resulted in empty text")
            return jsonify({
                'error': 'Audio transcription resulted in empty text. Please ensure your recording is clear and try again.',
                'validation_errors': ['No text was transcribed from the audio']
            }), 400
        
        # Sanitize transcribed text to prevent prompt injection
        sanitized_transcription = sanitize_user_input(transcribed_text)
        if not sanitized_transcription or len(sanitized_transcription.strip()) < 10:
            return jsonify({
                'error': 'Transcribed content is invalid or too short after sanitization',
                'validation_errors': ['Content does not meet security requirements']
            }), 400
        
        # Validate sanitized content
        validation_result = validate_text_content(sanitized_transcription)
        
        # Update result with sanitized text
        if sanitized_transcription != transcribed_text:
            result['original_text'] = sanitized_transcription
            logger.info("Transcribed text was sanitized for security")
        
        if not validation_result['valid']:
            logger.warning(f"Transcribed text validation failed: {validation_result['errors']}")
            return jsonify({
                'error': 'Audio content validation failed',
                'validation_errors': validation_result['errors'],
                'transcribed_text': transcribed_text
            }), 400
        
        # Add project info to result
        result['project_info'] = project_info
        result['source_type'] = 'audio_recording'
        
        logger.info(f"Audio processing completed successfully")
        return jsonify(result)
            
    except json.JSONDecodeError as e:
        logger.error(f"JSON decode error: {str(e)}")
        return jsonify({'error': f'Invalid project_info format: {str(e)}'}), 400
    except Exception as e:
        logger.error(f"Error processing audio requirement: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to process audio: {str(e)}'}), 500
    finally:
        # Clean up temporary file
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
                logger.info(f"Cleaned up temp file: {temp_file_path}")
            except Exception as cleanup_error:
                logger.warning(f"Failed to cleanup temp file {temp_file_path}: {cleanup_error}")

@app.route('/api/transcribe-audio', methods=['POST'])
def transcribe_audio_only():
    """Transcribe audio without full processing - for live transcription during recording"""
    temp_file_path = None
    try:
        audio_file = request.files.get('audio')

        if not audio_file:
            return jsonify({'error': 'No audio file provided'}), 400

        # Note: some browsers/clients don't set content_length on multipart parts,
        # so rely on the saved temp file size check below instead of rejecting here.
        logger.info(f"Transcription request: {audio_file.filename}, Size: {getattr(audio_file, 'content_length', None)} bytes")
        
        # Determine file extension
        original_filename = audio_file.filename or 'recording'
        file_ext = None
        
        if '.' in original_filename:
            file_ext = original_filename.rsplit('.', 1)[1].lower()
        
        content_type = audio_file.content_type or ''
        if not file_ext or file_ext not in ['wav', 'webm', 'mp4', 'm4a', 'ogg', 'mp3', 'flac']:
            if 'webm' in content_type or 'webm' in original_filename.lower():
                file_ext = 'webm'
            elif 'mp4' in content_type or 'm4a' in content_type:
                file_ext = 'm4a'
            elif 'ogg' in content_type:
                file_ext = 'ogg'
            elif 'wav' in content_type:
                file_ext = 'wav'
            else:
                file_ext = 'webm'
        
        # Create temporary file
        suffix = f'.{file_ext}'
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temp_file:
            audio_file.save(temp_file.name)
            temp_file_path = temp_file.name
        
        # Reject empty files early
        if not os.path.exists(temp_file_path) or os.path.getsize(temp_file_path) == 0:
            logger.error(f"Transcription file is empty: {temp_file_path}")
            return jsonify({'error': 'Audio file is empty. Please record again.'}), 400

        # Reject very short audio (< 0.5s)
        try:
            duration = orchestrator.processor._get_audio_duration(temp_file_path)
            if duration <= 0.5:
                logger.error(f"Transcription file too short ({duration}s): {temp_file_path}")
                return jsonify({'error': 'Audio is too short. Please record at least 1 second.'}), 400
        except Exception as dur_err:
            logger.warning(f"Failed to measure duration for {temp_file_path}: {dur_err}")
        
        # Transcribe using orchestrator's processor
        try:
            # Ensure whisper is loaded
            if not getattr(orchestrator.processor, 'models_loaded', False):
                orchestrator.processor._load_models()
            transcription = orchestrator.processor._transcribe_audio(temp_file_path)
            logger.info(f"Transcription successful, length: {len(transcription)} characters")
            
            return jsonify({
                'transcription': transcription,
                'success': True
            })
        except Exception as transcribe_error:
            logger.error(f"Transcription error: {str(transcribe_error)}")
            return jsonify({
                'error': f'Transcription failed: {str(transcribe_error)}',
                'transcription': ''
            }), 500
        
    except Exception as e:
        logger.error(f"Error transcribing audio: {str(e)}", exc_info=True)
        return jsonify({'error': f'Transcription failed: {str(e)}', 'transcription': ''}), 500
    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
            except Exception:
                pass

@app.route('/api/process-batch', methods=['POST'])
def process_batch_requirements():
    """Process multiple requirements from uploaded files"""
    try:
        files = request.files.getlist('files')
        project_info_str = request.form.get('project_info', '{}')
        
        if not files:
            return jsonify({'error': 'No files provided'}), 400
        
        project_info = json.loads(project_info_str)
        
        # Create temporary directory for processing
        with tempfile.TemporaryDirectory() as temp_dir:
            results = []
            validation_errors_list = []
            
            for file in files:
                if file.filename == '':
                    continue
                
                # Save file to temp directory
                file_path = os.path.join(temp_dir, file.filename)
                file.save(file_path)

                # Reject zero-size after save
                if not os.path.exists(file_path) or os.path.getsize(file_path) == 0:
                    validation_errors_list.append({
                        'file': file.filename,
                        'errors': ['File is empty. Please upload a valid file.']
                    })
                    continue
                
                # Determine input type and read content for validation
                lower_name = file.filename.lower()
                if lower_name.endswith(('.wav', '.mp3', '.m4a', '.flac')):
                    input_data = {'type': 'audio', 'file_path': file_path}
                    # Process first to get transcription
                    result = orchestrator.process_single_requirement(input_data)
                    content_to_validate = result.get('original_text', '')
                elif lower_name.endswith('.txt'):
                    # Read plain text file
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()

                    # Sanitize file content to prevent prompt injection
                    sanitized_content = sanitize_user_input(content)
                    if not sanitized_content or len(sanitized_content.strip()) < 10:
                        validation_errors_list.append({
                            'file': file.filename,
                            'errors': ['Content is invalid or too short after sanitization']
                        })
                        continue
                    
                    content_to_validate = sanitized_content
                    input_data = {'type': 'text', 'content': sanitized_content}
                    # Process the requirement
                    result = orchestrator.process_single_requirement(input_data)
                elif lower_name.endswith('.pdf'):
                    content = _extract_text_from_pdf(file_path) or ""
                    sanitized_content = sanitize_user_input(content)
                    if not sanitized_content or len(sanitized_content.strip()) < 10:
                        validation_errors_list.append({
                            'file': file.filename,
                            'errors': ['PDF content is invalid or too short after extraction/sanitization']
                        })
                        continue
                    content_to_validate = sanitized_content
                    input_data = {'type': 'text', 'content': sanitized_content}
                    result = orchestrator.process_single_requirement(input_data)
                elif lower_name.endswith('.docx'):
                    try:
                        doc = Document(file_path)
                        content = "\n".join(p.text for p in doc.paragraphs if str(p.text).strip())
                    except Exception:
                        content = ""
                    sanitized_content = sanitize_user_input(content)
                    if not sanitized_content or len(sanitized_content.strip()) < 10:
                        validation_errors_list.append({
                            'file': file.filename,
                            'errors': ['DOCX content is invalid or too short after extraction/sanitization']
                        })
                        continue
                    content_to_validate = sanitized_content
                    input_data = {'type': 'text', 'content': sanitized_content}
                    result = orchestrator.process_single_requirement(input_data)
                else:
                    validation_errors_list.append({
                        'file': file.filename,
                        'errors': ['Unsupported file type. Allowed: .txt, .docx, .pdf']
                    })
                    continue
                
                # Validate the content
                validation_result = validate_text_content(content_to_validate)
                
                if not validation_result['valid']:
                    validation_errors_list.append({
                        'file': file.filename,
                        'errors': validation_result['errors']
                    })
                
                result['source_file'] = file.filename
                result['validation'] = validation_result
                results.append(result)
            
            # If any validation errors, return them
            if validation_errors_list:
                return jsonify({
                    'error': 'File content validation failed',
                    'validation_errors': validation_errors_list,
                    'details': 'One or more files do not meet the requirements'
                }), 400
            
            # Add project info to results
            batch_result = {
                'results': results,
                'project_info': project_info,
                'timestamp': datetime.now().isoformat(),
                'total_files': len(results),
                'status': 'completed'
            }
            
            return jsonify(batch_result)
        
    except Exception as e:
        logger.error(f"Error processing batch requirements: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/generate-srs', methods=['POST'])
def generate_srs():
    """
    Generate SRS (Software Requirements Specification) document from processed requirements.
    
    Accepts processed requirement results and project information, then generates
    a complete IEEE 830-1998 compliant SRS document using the model-based generator.
    
    Expected request body:
        - results: List of processed requirement objects
        - project_info: Dictionary containing project metadata (title, author, version)
    
    Returns:
        JSON response containing the generated SRS document with:
            - document_id: Unique identifier for the document
            - title: Document title
            - version: Document version
            - date: Generation date
            - author: Document author
            - sections: Parsed SRS sections (introduction, overall_description, specific_requirements)
            - raw_text: Full raw text output from the model
    """
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        results = data.get('results')
        project_info = data.get('project_info', {})
        
        if not results:
            return jsonify({'error': 'No results provided'}), 400
        
        # Ensure results is a list
        if not isinstance(results, list):
            results = [results]

        srs_dict = _build_srs_dict_from_results(results, project_info)
        return jsonify(srs_dict)

    except RequirementSecurityError as sec:
        return jsonify(sec.payload), sec.status

    except Exception as e:
        logger.error(f"Error generating SRS: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/generate-srs-stream', methods=['POST'])
def generate_srs_stream():
    """
    Stream SRS generation over Server-Sent Events (SSE).

    Emits JSON lines in SSE `data:` frames:
      - {"type": "delta", "text": "..."} — append to the live preview (token/chunk-level or chunked RAG text).
      - {"type": "done", "srs": {...}} — same payload shape as POST /api/generate-srs.
      - {"type": "error", "message": "..."} on failure.

    Streaming reduces perceived latency: the client can render the first tokens before the full document exists.
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        results = data.get('results')
        project_info = data.get('project_info', {})
        if not results:
            return jsonify({'error': 'No results provided'}), 400
        if not isinstance(results, list):
            results = [results]

        results = sanitize_requirement_results(results)
        project_info = sanitize_project_info(project_info)
        inj = reject_payload_if_prompt_injection_in_results(results)
        if inj:
            raise RequirementSecurityError(inj, 403)

        logger.info(f"SRS stream: generating with {len(results)} result(s)")

        resolved_mode = str(os.environ.get("SRS_GENERATION_MODE", "model")).strip().lower()
        if resolved_mode not in {"rag", "model"}:
            resolved_mode = "model"

        @stream_with_context
        def event_stream():
            try:
                if resolved_mode == "rag":
                    # RAG returns a full document; we still chunk `raw_text` so the UI can reveal it progressively.
                    srs = _generate_srs_document(results, project_info, mode_override="rag")
                    raw = (getattr(srs, "raw_text", None) or (srs.sections or {}).get("_raw_text") or "")
                    step = 320
                    for i in range(0, len(raw), step):
                        yield _sse_srs_event({"type": "delta", "text": raw[i : i + step]})
                    out = _srs_document_to_api_dict(srs, results, project_info, "rag")
                    gm = dict(out.get("generation_meta") or {})
                    gm["streamed"] = True
                    out["generation_meta"] = gm
                    yield _sse_srs_event({"type": "done", "srs": out})
                    return

                token = (os.environ.get("REPLICATE_API_TOKEN") or "").strip()
                if not token:
                    yield _sse_srs_event(
                        {
                            "type": "error",
                            "message": "REPLICATE_API_TOKEN is not configured; streaming requires the model backend.",
                        }
                    )
                    return

                model_gen = SRSModelGenerator()
                req_text, prompt, chunk_it = model_gen.stream_srs_text_chunks(results, project_info)
                parts = []
                for piece in chunk_it:
                    parts.append(piece)
                    yield _sse_srs_event({"type": "delta", "text": piece})
                full_raw = "".join(parts)
                srs_doc = model_gen._document_from_replicate_output(full_raw, prompt, req_text, project_info)
                out = _srs_document_to_api_dict(srs_doc, results, project_info, "model")
                gm = dict(out.get("generation_meta") or {})
                gm["streamed"] = True
                out["generation_meta"] = gm
                yield _sse_srs_event({"type": "done", "srs": out})
            except Exception as e:
                logger.exception("SRS stream failed inside generator")
                yield _sse_srs_event({"type": "error", "message": str(e)})

        return Response(
            event_stream(),
            mimetype="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "X-Accel-Buffering": "no",
            },
        )

    except RequirementSecurityError as sec:
        return jsonify(sec.payload), sec.status
    except Exception as e:
        logger.error(f"Error starting SRS stream: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/generate-srs-compare', methods=['POST'])
def generate_srs_compare():
    """
    Testing endpoint: generate two SRS variants from same input for side-by-side comparison.
    Returns both model-only and rag outputs with simple delta indicators.
    """
    try:
        data = request.get_json() or {}
        results = data.get('results')
        project_info = data.get('project_info', {})

        if not results:
            return jsonify({'error': 'No results provided'}), 400
        if not isinstance(results, list):
            results = [results]

        model_output = _build_srs_dict_from_results(results, project_info, mode_override="model")
        rag_output = _build_srs_dict_from_results(results, project_info, mode_override="rag")

        model_hall = model_output.get("hallucination_analysis", {}) or {}
        rag_hall = rag_output.get("hallucination_analysis", {}) or {}
        model_raw = str(model_output.get("raw_text", "") or "")
        rag_raw = str(rag_output.get("raw_text", "") or "")
        model_flags = list(model_hall.get("flagged_sections", []) or [])
        rag_flags = list(rag_hall.get("flagged_sections", []) or [])

        comparison = {
            "raw_length": {"model": len(model_raw), "rag": len(rag_raw), "delta": len(rag_raw) - len(model_raw)},
            "confidence_score": {
                "model": float(model_hall.get("confidence_score", 0.0) or 0.0),
                "rag": float(rag_hall.get("confidence_score", 0.0) or 0.0),
                "delta": float(rag_hall.get("confidence_score", 0.0) or 0.0) - float(model_hall.get("confidence_score", 0.0) or 0.0),
            },
            "flagged_sections_count": {
                "model": len(model_flags),
                "rag": len(rag_flags),
                "delta": len(rag_flags) - len(model_flags),
            },
        }

        return jsonify({
            "model_output": model_output,
            "rag_output": rag_output,
            "comparison": comparison,
        })
    except RequirementSecurityError as sec:
        return jsonify(sec.payload), sec.status
    except Exception as e:
        logger.error(f"Error generating SRS comparison: {str(e)}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/rtm-analyze', methods=['POST'])
def rtm_analyze():
    """Build Requirements Traceability Matrix report from SRS + use case artifacts."""
    try:
        data = request.get_json() or {}
        srs_data = data.get("srs_data") or {}
        use_case_data = data.get("use_case_data") or {}
        if not isinstance(srs_data, dict) or not srs_data:
            return jsonify({"error": "srs_data is required"}), 400
        report = _build_rtm_report(srs_data, use_case_data if isinstance(use_case_data, dict) else {})
        return jsonify(report)
    except Exception as e:
        logger.error("rtm-analyze failed: %s", e, exc_info=True)
        return jsonify({"error": str(e), "summary": {}, "rows": []}), 500


@app.route('/api/generate-usecases', methods=['POST'])
def generate_usecases():
    """Generate textual use cases and optional use case diagram from SRS sections."""
    try:
        data = request.get_json() or {}
        sections = data.get('sections') or {}
        document_id = data.get('document_id', f"SRS-{datetime.now().strftime('%Y%m%d-%H%M%S')}")
        title = data.get('title', 'System')

        if not sections:
            return jsonify({'error': 'No SRS sections provided'}), 400

        output_dir = Path('data/output')
        output_dir.mkdir(parents=True, exist_ok=True)
        safe_id = re.sub(r'[^A-Za-z0-9_.-]', '_', document_id)

        textual_output = output_dir / f"textual_usecases_{safe_id}.txt"
        textual_gen = TextualUseCaseGenerator()
        textual_result = textual_gen.generate_and_save(
            srs_sections=sections,
            output_path=str(textual_output),
        )

        diagram_gen = UseCaseDiagramGenerator()
        diagram_bundle = diagram_gen.generate_both_layouts_and_render(
            textual_use_cases=textual_result['text'],
            system_name=title,
            output_dir=str(output_dir),
            output_name=f"usecase_{safe_id}",
        )
        dv = diagram_bundle.get('vertical') or {}
        dh = diagram_bundle.get('horizontal') or {}

        def _png_b64(abs_path: str) -> str:
            if abs_path and Path(abs_path).exists():
                with open(abs_path, 'rb') as img_file:
                    return base64.b64encode(img_file.read()).decode('utf-8')
            return ''

        diagram_base64_v = _png_b64(dv.get('diagram_file', ''))
        diagram_base64_h = _png_b64(dh.get('diagram_file', ''))
        diagram_base64 = diagram_base64_v or diagram_base64_h
        diagram_abs_path = dv.get('diagram_file') or dh.get('diagram_file') or ''
        diagram_rel_path = ""
        if diagram_abs_path and Path(diagram_abs_path).exists():
            diagram_rel_path = str(Path(diagram_abs_path).resolve().relative_to(Path.cwd().resolve())).replace("\\", "/")

        msg_parts = []
        if dv.get('status') != 'rendered' and dv.get('message'):
            msg_parts.append(f"Vertical: {dv.get('message')}")
        if dh.get('status') != 'rendered' and dh.get('message'):
            msg_parts.append(f"Horizontal: {dh.get('message')}")
        combined_msg = " ".join(msg_parts) if msg_parts else ""
        log_v = (dv.get('plantuml_log') or '').strip()
        log_h = (dh.get('plantuml_log') or '').strip()
        combined_log = "\n---\n".join(x for x in (log_v, log_h) if x)

        diag_status = 'rendered' if (diagram_base64_v or diagram_base64_h) else 'saved_only'

        return jsonify({
            'textual_usecases': {
                'use_cases': textual_result.get('use_cases', []),
                'text': textual_result.get('text', ''),
                'output_file': textual_result.get('output_file', ''),
            },
            'diagram': {
                'status': diag_status,
                'plantuml_code': dv.get('plantuml_code', ''),
                'plantuml_code_vertical': dv.get('plantuml_code', ''),
                'plantuml_code_horizontal': dh.get('plantuml_code', ''),
                'puml_file': dv.get('puml_file', ''),
                'puml_file_horizontal': dh.get('puml_file', ''),
                'diagram_file': diagram_abs_path,
                'diagram_rel_path': diagram_rel_path,
                'diagram_base64': diagram_base64,
                'diagram_base64_vertical': diagram_base64_v,
                'diagram_base64_horizontal': diagram_base64_h,
                'message': combined_msg or dv.get('message', '') or dh.get('message', ''),
                'plantuml_log': combined_log,
            }
        })
    except Exception as e:
        logger.error(f"Error generating use cases: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/download-srs/<format>', methods=['POST'])
def download_srs(format):
    """Download SRS document in specified format"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        srs_data = data.get('srs_data')
        
        if not srs_data:
            return jsonify({'error': 'No SRS data provided'}), 400
        
        # Create temporary file
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if format.lower() == 'json':
            filename = f"srs_{timestamp}.json"
            with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False) as f:
                json.dump(srs_data, f, indent=2, ensure_ascii=False)
                temp_file = f.name
        elif format.lower() == 'html':
            filename = f"srs_{timestamp}.html"
            # Generate HTML content
            html_content = generate_html_content(srs_data)
            with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as f:
                f.write(html_content)
                temp_file = f.name
        else:
            return jsonify({'error': 'Unsupported format'}), 400
        
        return send_file(
            temp_file,
            as_attachment=True,
            download_name=filename,
            mimetype='application/octet-stream'
        )
        
    except Exception as e:
        logger.error(f"Error downloading SRS: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/download-textual-usecases-pdf', methods=['POST'])
def download_textual_usecases_pdf():
    """Render textual use cases as PDF (or HTML fallback)."""
    try:
        data = request.get_json() or {}
        text_payload = str(data.get('text', '')).strip()
        title = str(data.get('title', 'Textual Use Cases')).strip() or 'Textual Use Cases'
        if not text_payload:
            return jsonify({'error': 'No textual use case content provided'}), 400

        import html
        def _textual_usecase_download_name(src_title: str) -> str:
            raw = str(src_title or "").strip()
            # UI title may include a suffix like " - Textual Use Cases".
            raw = re.sub(r'\s*-\s*Textual\s+Use\s+Cases\s*$', '', raw, flags=re.IGNORECASE).strip()
            if not raw:
                raw = "Project"
            safe = re.sub(r'[\\/:*?"<>|]+', '', raw)
            safe = re.sub(r'\s+', '_', safe).strip('._ ')
            if not safe:
                safe = "Project"
            return f"TextualUseCase_{safe}.pdf"

        def _format_textual_usecases_html(text: str):
            """
            Parse textual use cases into structured cards.
            Returns: (body_html, toc_items)
            """
            payload = (text or "").replace('\r\n', '\n')
            lines = [ln.strip() for ln in payload.split('\n')]

            clean_lines = []
            for i, ln in enumerate(lines):
                if not ln:
                    clean_lines.append('')
                    continue
                # Skip duplicated document title/banner lines in payload.
                if i == 0 and re.search(r'textual\s+use\s+cases', ln, flags=re.IGNORECASE):
                    continue
                if ln == '-':
                    continue
                ln = re.sub(r'\s+-\s*$', '', ln).strip()
                clean_lines.append(ln)

            cases = []
            current = None
            field_keys = {
                'Primary Actor',
                'Stakeholders and Interests',
                'Preconditions',
                'Postconditions',
                'Main Success Scenario',
                'Extensions',
                'Special Requirements',
                'Frequency of Occurrence',
                'Assumptions',
            }

            def flush_case():
                nonlocal current
                if current and current.get('title'):
                    cases.append(current)
                current = None

            active_field = None
            for idx, ln in enumerate(clean_lines):
                if not ln:
                    continue

                m = re.match(r'^([^:]{2,80})\s*:\s*(.*)$', ln)
                is_bullet = ln.startswith('- ') or ln.startswith('* ')
                is_step = bool(re.match(r'^\d+\.\s+', ln))
                next_non_empty = ''
                for j in range(idx + 1, len(clean_lines)):
                    if clean_lines[j]:
                        next_non_empty = clean_lines[j]
                        break
                next_kv = re.match(r'^([^:]{2,80})\s*:\s*(.*)$', next_non_empty) if next_non_empty else None

                # Treat non key-value + non-bullet lines as use case titles.
                likely_title = (
                    not m and not is_bullet and not is_step and
                    (
                        current is None or
                        (next_non_empty == '-' or (next_kv and (next_kv.group(1).strip() in field_keys or re.match(r'^Use Case Name$', next_kv.group(1).strip(), flags=re.IGNORECASE))))
                    )
                )
                if likely_title:
                    flush_case()
                    current = {'title': ln, 'fields': {}, 'bullets': []}
                    active_field = None
                    continue

                if current is None:
                    current = {'title': 'Use Case', 'fields': {}, 'bullets': []}

                if m:
                    key = m.group(1).strip()
                    value = (m.group(2) or '').strip()
                    if re.match(r'^Use Case Name$', key, flags=re.IGNORECASE):
                        # Start a new case when a new Use Case Name appears.
                        has_content = bool(current.get('fields')) or bool(current.get('bullets')) or (
                            current.get('title') and current.get('title') != 'Use Case'
                        )
                        if has_content:
                            flush_case()
                            current = {'title': 'Use Case', 'fields': {}, 'bullets': []}
                        if value:
                            current['title'] = value
                        active_field = None
                    else:
                        if key in field_keys or key:
                            current['fields'][key] = value
                            active_field = key
                    continue

                if ln.startswith('- ') or ln.startswith('* '):
                    current['bullets'].append(ln[2:].strip())
                    active_field = None
                else:
                    # Continuation line for wrapped field values (prevents truncation of first use case).
                    if active_field:
                        prev = current['fields'].get(active_field, '')
                        current['fields'][active_field] = (f'{prev} {ln}'.strip() if prev else ln)
                    else:
                        current['bullets'].append(ln)

            flush_case()

            if not cases:
                return f'<p class="muted">{html.escape(text)}</p>', []

            toc_items = []
            rendered = []

            for idx, case in enumerate(cases, start=1):
                title_text = f'UC-{idx:02d}: {case["title"]}'
                anchor = f'uc-{idx:02d}'
                toc_items.append((title_text, anchor))

                def row(label: str) -> str:
                    raw_val = case['fields'].get(label, '')
                    val = html.escape(raw_val) if raw_val else '<span class="muted">Not specified</span>'
                    return f'<tr><th>{html.escape(label)}</th><td>{val}</td></tr>'

                extra_html = ''
                if case['bullets']:
                    extra_html = (
                        '<div class="extra"><p class="extra-title">Additional Notes</p><ul class="list">'
                        + ''.join(f'<li>{html.escape(x)}</li>' for x in case['bullets'] if x)
                        + '</ul></div>'
                    )

                if idx > 1:
                    rendered.append('<pdf:nextpage />')
                rendered.append(
                    f'<section class="uc-card"><a name="{anchor}"></a>'
                    f'<h2 class="uc-title">{html.escape(title_text)}</h2>'
                    f'<table class="uc-table">'
                    f'{row("Primary Actor")}'
                    f'{row("Stakeholders and Interests")}'
                    f'{row("Preconditions")}'
                    f'{row("Postconditions")}'
                    f'{row("Main Success Scenario")}'
                    f'{row("Extensions")}'
                    f'{row("Special Requirements")}'
                    f'{row("Frequency of Occurrence")}'
                    f'{row("Assumptions")}'
                    f'</table>{extra_html}</section>'
                )

            return ''.join(rendered), toc_items

        body_html, toc_items = _format_textual_usecases_html(text_payload)
        toc_html = ''.join(
            f'<li><a href="#{html.escape(anchor)}">{html.escape(label)}</a></li>'
            for label, anchor in toc_items
        )
        html_doc = f"""<!DOCTYPE html>
<html>
<head>
  <meta charset="UTF-8">
  <title>{html.escape(title)}</title>
  <style>
    @page {{
      size: A4;
      margin: 24mm 18mm 20mm 18mm;
      @bottom-right {{
        content: "Page " counter(page) " of " counter(pages);
        font-size: 9pt;
        color: #475569;
      }}
    }}
    body {{
      font-family: 'Times New Roman', Times, serif;
      color: #0f172a;
      line-height: 1.6;
      font-size: 11.5pt;
      margin: 0;
    }}
    .cover {{
      border: 1px solid #94a3b8;
      border-radius: 8px;
      padding: 20px;
      min-height: 235mm;
      box-sizing: border-box;
      page-break-after: always;
    }}
    .cover-kicker {{
      margin: 0 0 8px;
      font-size: 10pt;
      letter-spacing: 1pt;
      text-transform: uppercase;
      color: #475569;
      text-align: center;
    }}
    .cover-title {{
      margin: 0 0 12px;
      font-size: 30pt;
      font-weight: 800;
      text-align: center;
    }}
    .cover-sub {{
      margin: 0 0 16px;
      font-size: 12pt;
      text-align: center;
      color: #334155;
      font-weight: 600;
    }}
    .meta {{
      border: 1px solid #cbd5e1;
      border-radius: 6px;
      padding: 10px 12px;
      background: #f8fafc;
      margin-bottom: 14px;
    }}
    .meta p {{ margin: 3px 0; }}
    .toc {{
      border: 1px solid #cbd5e1;
      border-radius: 8px;
      padding: 14px;
      min-height: 235mm;
      box-sizing: border-box;
      page-break-after: always;
    }}
    .toc h2 {{
      margin: 0 0 10px;
      font-size: 18pt;
      font-weight: 800;
      border-bottom: 1px solid #dbe3ee;
      padding-bottom: 4px;
    }}
    .toc ul {{ margin: 0; padding-left: 18px; }}
    .toc li {{ margin: 4px 0; }}
    .toc a {{ color: #0f172a; text-decoration: none; }}
    .toc a:hover {{ text-decoration: underline; }}
    .content {{ margin-top: 0; }}
    .uc-table {{
      width: 100%;
      border-collapse: collapse;
      margin: 6px 0 10px;
      font-size: 10.8pt;
    }}
    .uc-table th, .uc-table td {{
      border: 1px solid #cbd5e1;
      padding: 6px 7px;
      text-align: left;
      vertical-align: top;
    }}
    .uc-table th {{
      width: 28%;
      background: #f8fafc;
      font-weight: 700;
    }}
    .uc-card {{
      border: 1px solid #d1d5db;
      border-radius: 8px;
      padding: 12px 14px;
      margin: 0 0 12px;
      background: #fff;
      page-break-inside: avoid;
      page-break-before: always;
    }}
    .uc-title {{
      margin: 0 0 8px;
      font-size: 15pt;
      font-weight: 800;
      border-bottom: 1px solid #e2e8f0;
      padding-bottom: 5px;
      background: #f8fbff;
      border-left: 4px solid #9ab6db;
      padding-left: 8px;
      border-radius: 3px;
    }}
    .extra-title {{
      margin: 0 0 6px;
      font-size: 11pt;
      font-weight: 700;
    }}
    .list {{ margin: 4px 0 6px 20px; }}
    .list li {{ margin: 4px 0; }}
    .muted {{ color: #64748b; font-style: italic; }}
    .doc-footer {{
      margin-top: 18px;
      border-top: 1px solid #cbd5e1;
      padding-top: 8px;
      color: #64748b;
      font-size: 10pt;
    }}
  </style>
</head>
<body>
  <section class="cover">
    <p class="cover-kicker">Textual Use Cases</p>
    <h1 class="cover-title">{html.escape(title)}</h1>
    <p class="cover-sub">Req2Design · Structured Use Case Specification</p>
    <div class="meta">
      <p><strong>Document Type:</strong> Textual Use Cases</p>
      <p><strong>Generated On:</strong> {html.escape(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}</p>
      <p><strong>Total Use Cases:</strong> {len(toc_items)}</p>
    </div>
  </section>
  <pdf:nextpage />

  <section class="toc">
    <h2>Table of Contents</h2>
    <ul>{toc_html}</ul>
  </section>
  <pdf:nextpage />

  <div class="content">{body_html}</div>
  <div class="doc-footer"><strong>End of document.</strong> Generated by Req2Design textual use case exporter.</div>
</body>
</html>"""
        out_path = save_pdf_or_html(html_doc, "textual_usecases.pdf")
        is_pdf = out_path.endswith('.pdf')
        _resp = send_file(
            out_path,
            as_attachment=True,
            download_name=_textual_usecase_download_name(title),
            mimetype='application/pdf' if is_pdf else 'text/html; charset=utf-8',
        )
        _resp.headers['X-Export-Format'] = 'pdf' if is_pdf else 'html'
        return _resp
    except Exception as e:
        logger.error(f"Error generating textual use cases PDF: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/download-usecase-diagram-pdf', methods=['POST'])
def download_usecase_diagram_pdf():
    """Render use case diagram image as PDF (or HTML fallback)."""
    try:
        data = request.get_json() or {}
        image_base64 = str(data.get('diagram_base64', '')).strip()
        title = str(data.get('title', 'Use Case Diagram')).strip() or 'Use Case Diagram'
        if not image_base64:
            return jsonify({'error': 'No diagram image provided'}), 400

        import html
        html_doc = f"""<!DOCTYPE html>
<html>
<head>
  <meta charset="UTF-8">
  <title>{html.escape(title)}</title>
  <style>
    body {{ font-family: Arial, sans-serif; margin: 28px; color: #111827; }}
    h1 {{ margin-bottom: 16px; }}
    .wrap {{ text-align: center; }}
    img {{ max-width: 100%; height: auto; border: 1px solid #d1d5db; border-radius: 6px; }}
  </style>
</head>
<body>
  <h1>{html.escape(title)}</h1>
  <div class="wrap">
    <img src="data:image/png;base64,{image_base64}" alt="Use Case Diagram" />
  </div>
</body>
</html>"""
        out_path = save_pdf_or_html(html_doc, "usecase_diagram.pdf")
        is_pdf = out_path.endswith('.pdf')
        _resp = send_file(
            out_path,
            as_attachment=True,
            download_name=os.path.basename(out_path),
            mimetype='application/pdf' if is_pdf else 'text/html; charset=utf-8',
        )
        _resp.headers['X-Export-Format'] = 'pdf' if is_pdf else 'html'
        return _resp
    except Exception as e:
        logger.error(f"Error generating use case diagram PDF: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/generate-srs-pdf', methods=['POST'])
def generate_srs_pdf():
    """
    Generate SRS document as PDF (or HTML fallback) and return as downloadable file.
    
    This endpoint generates a PDF file from SRS data. If PDF generation fails
    (e.g., weasyprint not installed), it falls back to HTML format.
    
    The function prioritizes raw_text if available for full document fidelity,
    otherwise uses parsed sections to reconstruct the document.
    
    Expected request body:
        - document_id: Unique identifier for the document
        - title: Document title
        - version: Document version
        - date: Document date
        - author: Document author
        - raw_text: Full raw text from model (preferred)
        - sections: Parsed SRS sections (fallback if raw_text unavailable)
        - results: Processed requirements (only used if sections are empty)
        - project_info: Project metadata
    
    Returns:
        File download response (PDF or HTML) with appropriate MIME type
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400

        def _project_pdf_download_name(title: str, document_id: str) -> str:
            """
            Build a safe downloadable filename from project title.
            """
            raw = str(title or "").strip()
            if not raw:
                raw = f"{document_id or 'document'}"
            safe = re.sub(r'[\\/:*?"<>|]+', '', raw)
            safe = re.sub(r'\s+', '_', safe).strip('._ ')
            if not safe:
                safe = f"{document_id or 'document'}"
            return f"SRS_{safe}.pdf"

        # Helper function to check if sections are empty
        def _is_empty_sections(sec: dict) -> bool:
            """
            Check if SRS sections dictionary is effectively empty.
            
            Args:
                sec: Dictionary containing SRS sections
            
            Returns:
                True if sections contain no meaningful content, False otherwise
            """
            if not isinstance(sec, dict):
                return True
            intro = sec.get('introduction') or {}
            overall = sec.get('overall_description') or {}
            has_intro = any(bool(str(intro.get(k, '')).strip()) for k in ['purpose','scope','overview']) or bool(intro.get('definitions'))
            has_overall = any(bool(overall.get(k)) for k in ['product_functions','user_characteristics','constraints','assumptions','dependencies']) or bool(str(overall.get('product_perspective','')).strip())
            return not (has_intro or has_overall)

        # ALWAYS prioritize raw_text if available - this is the exact model output
        raw_text = data.get('raw_text')
        logger.info(f"PDF generation request - raw_text present: {raw_text is not None}, length: {len(raw_text) if raw_text else 0}")
        
        if raw_text and len(raw_text.strip()) > 50:  # Lowered threshold to 50 chars
            logger.info(f"Using provided raw_text for full SRS document generation (length: {len(raw_text)} chars)")
            doc_id = data.get('document_id', f"SRS-{datetime.now().strftime('%Y%m%d-%H%M%S')}")
            title = data.get('title', 'Software Requirements Specification')
            # Convert raw text to HTML with proper formatting - NO metadata, NO parsing, just raw text
            html = _convert_raw_text_to_html(
                raw_text,
                doc_id,
                title,
                data.get('version', '1.0'),
                data.get('date', datetime.now().strftime('%Y-%m-%d')),
                data.get('author', 'System')
            )
            out_path = save_pdf_or_html(html, f"srs_{doc_id}.pdf")
            # Determine MIME type and file extension based on actual output
            is_pdf = out_path.endswith('.pdf')
            mimetype = 'application/pdf' if is_pdf else 'text/html; charset=utf-8'
            download_name = _project_pdf_download_name(title, doc_id)
            
            _resp = send_file(
                out_path,
                as_attachment=True,
                download_name=download_name,
                mimetype=mimetype,
            )
            _resp.headers['X-Export-Format'] = 'pdf' if is_pdf else 'html'
            return _resp
        
        # If raw_text is missing or too short, log warning
        if not raw_text:
            logger.warning("raw_text not provided in request - will fall back to parsed sections")
        elif len(raw_text.strip()) <= 50:
            logger.warning(f"raw_text too short ({len(raw_text.strip())} chars) - will fall back to parsed sections")

        sections = data.get('sections') or data.get('srs_sections')
        
        # Only use parsed sections if raw_text is NOT available
        # Check if sections are already provided (from /api/generate-srs)
        if sections and not _is_empty_sections(sections):
            logger.info("Using provided SRS sections (no regeneration needed)")
            # Use provided sections - no regeneration
            srs_data = {
                'document_id': data.get('document_id', f"SRS-{datetime.now().strftime('%Y%m%d-%H%M%S')}"),
                'title': data.get('title', 'Software Requirements Specification'),
                'version': data.get('version', '1.0'),
                'date': data.get('date', datetime.now().strftime('%Y-%m-%d')),
                'author': data.get('author', 'System'),
                'sections': sections,
            }
            html = render_html(srs_data)
            out_path = save_pdf_or_html(html, f"srs_{srs_data['document_id']}.pdf")
            is_pdf = out_path.endswith('.pdf')
            mimetype = 'application/pdf' if is_pdf else 'text/html; charset=utf-8'
            _resp = send_file(
                out_path,
                as_attachment=True,
                download_name=_project_pdf_download_name(srs_data.get('title'), srs_data.get('document_id')),
                mimetype=mimetype,
            )
            _resp.headers['X-Export-Format'] = 'pdf' if is_pdf else 'html'
            return _resp
        
        # Only generate if sections are truly empty AND results are provided
        # This should rarely happen if frontend calls /api/generate-srs first
        if data.get('results'):
            logger.warning("Sections empty but results provided - generating SRS (this may cause duplicate requests)")
            project_info = data.get('project_info', {})
            srs = _generate_srs_document(data['results'], project_info)
            # Use raw_text if available, otherwise use sections
            raw_text = getattr(srs, 'raw_text', None) or srs.sections.get('_raw_text')
            if raw_text and len(raw_text.strip()) > 100:
                logger.info("Using raw_text from SRS for full document generation")
                html = _convert_raw_text_to_html(
                    raw_text,
                    srs.document_id,
                    srs.title,
                    srs.version,
                    srs.date,
                    srs.author
                )
            else:
                html = render_html({
                    'document_id': srs.document_id,
                    'title': srs.title,
                    'version': srs.version,
                    'date': srs.date,
                    'author': srs.author,
                    'sections': srs.sections,
                })
            out_path = save_pdf_or_html(html, f"srs_{srs.document_id}.pdf")
            is_pdf = out_path.endswith('.pdf')
            mimetype = 'application/pdf' if is_pdf else 'text/html; charset=utf-8'
            _resp = send_file(
                out_path,
                as_attachment=True,
                download_name=_project_pdf_download_name(srs.title, srs.document_id),
                mimetype=mimetype,
            )
            _resp.headers['X-Export-Format'] = 'pdf' if is_pdf else 'html'
            return _resp
        
        # No sections and no results - return error
        return jsonify({'error': 'No SRS sections provided. Please generate SRS first using /api/generate-srs and pass the sections'}), 400
    except Exception as e:
        logger.error(f"Error generating SRS PDF: {str(e)}")
        return jsonify({'error': str(e)}), 500


def _build_srs_docx(document_id: str, title: str, version: str, date_value: str, author: str, raw_text: str, sections: dict):
    """
    Build a DOCX file path for SRS content.
    """
    def _normalize_for_export(text: str) -> str:
        x = (text or "").strip()
        # Fix cases like "1Introduction" / "1.1Purpose" by inserting space after numeric prefix.
        x = re.sub(r'(?m)^(\d+(?:\.\d+)*)([A-Za-z])', r'\1 \2', x)
        x = re.sub(r'(?m)^\s*-{3,}\s*$', '', x)
        x = re.sub(r'(?mi)^\s*===\s*[A-Z0-9 _/\-()]+(?:START|END)\s*===\s*$', '', x)
        x = re.sub(r'\s+(?=(INTRODUCTION|OVERALL DESCRIPTION|SPECIFIC REQUIREMENTS|SYSTEM FEATURES)\b)', '\n', x, flags=re.IGNORECASE)
        x = re.sub(r'\s+(?=\d+\.\s+[A-Z])', '\n', x)
        x = re.sub(r'\s+(?=\d+(?:\.\d+)+\s+[A-Z])', '\n', x)
        x = re.sub(r'\s+(?=FR-\d+\b)', '\n', x, flags=re.IGNORECASE)
        x = re.sub(r'\s+(?=NFR-\d+\b)', '\n', x, flags=re.IGNORECASE)
        x = re.sub(r'\s+-\s+', '\n- ', x)
        x = re.sub(r'\n{3,}', '\n\n', x)
        x = re.sub(r'(?m)^\s*-\s*$', '', x)
        x = re.sub(r'(?i)\bEnd of Document\..*$', 'End of Document.', x)

        intro_match = re.search(r'(?im)^\s*INTRODUCTION\s*$', x)
        if intro_match:
            x = x[intro_match.start():]
        end_marker = re.search(r'(?i)\bEnd of Document\.', x)
        if end_marker:
            x = x[:end_marker.end()]
        return _apply_ieee_numbering(x.strip())

    def _apply_ieee_numbering(text: str) -> str:
        lines = [ln.strip() for ln in text.splitlines()]
        out = []
        current_major = None
        current_sub = None

        major_map = {
            "introduction": (1, "Introduction"),
            "overall description": (2, "Overall Description"),
            "specific requirements": (3, "Specific Requirements"),
            "system features": (4, "System Features"),
        }
        sub_map = {
            1: {"purpose": (1, "Purpose"), "scope": (2, "Scope"), "definitions/acronyms": (3, "Definitions/Acronyms"), "references": (4, "References"), "overview": (5, "Overview")},
            2: {"product perspective": (1, "Product Perspective"), "product functions": (2, "Product Functions"), "user characteristics": (3, "User Characteristics"), "constraints": (4, "Constraints"), "assumptions/dependencies": (5, "Assumptions/Dependencies")},
            3: {"external interface requirements": (1, "External Interface Requirements"), "functional requirements": (2, "Functional Requirements"), "non-functional requirements": (3, "Non-Functional Requirements")},
        }
        subsub_map = {
            (3, 1): {"user interface": (1, "User Interface"), "hardware interface": (2, "Hardware Interface"), "software interface": (3, "Software Interface"), "communication interface": (4, "Communication Interface")},
            (3, 3): {"usability": (1, "Usability"), "reliability": (2, "Reliability"), "performance": (3, "Performance"), "portability": (4, "Portability")},
        }

        for line in lines:
            if not line:
                out.append("")
                continue
            if re.match(r'^\d+(?:\.\d+)*\s+', line):
                out.append(line)
                continue

            raw = line.strip().rstrip('.').rstrip(':')
            key = raw.lower()
            m = major_map.get(key)
            if m:
                current_major, title_txt = m
                current_sub = None
                out.append(f"{current_major}: {title_txt}")
                continue

            k = re.match(r'^([A-Za-z][A-Za-z0-9 /&()_-]{1,120})\s*:\s*(.*)$', line)
            if k and current_major in sub_map and k.group(1).strip().lower() in sub_map[current_major]:
                sub_idx, sub_label = sub_map[current_major][k.group(1).strip().lower()]
                current_sub = sub_idx
                tail = (k.group(2) or "").strip()
                if tail:
                    out.append(f"{current_major}.{sub_idx}: {sub_label}")
                    out.append(f"{sub_label}: {tail}")
                else:
                    out.append(f"{current_major}.{sub_idx}: {sub_label}")
                continue

            if k and (current_major, current_sub) in subsub_map and k.group(1).strip().lower() in subsub_map[(current_major, current_sub)]:
                subsub_idx, subsub_label = subsub_map[(current_major, current_sub)][k.group(1).strip().lower()]
                tail = (k.group(2) or "").strip()
                out.append(f"{current_major}.{current_sub}.{subsub_idx}: {subsub_label}" + (f": {tail}" if tail else ""))
                continue

            if current_major in sub_map and key in sub_map[current_major]:
                sub_idx, sub_label = sub_map[current_major][key]
                current_sub = sub_idx
                out.append(f"{current_major}.{sub_idx}: {sub_label}")
                continue
            if (current_major, current_sub) in subsub_map and key in subsub_map[(current_major, current_sub)]:
                subsub_idx, subsub_label = subsub_map[(current_major, current_sub)][key]
                out.append(f"{current_major}.{current_sub}.{subsub_idx}: {subsub_label}")
                continue

            out.append(line)

        return "\n".join(out).strip()

    def _heading_depth(line: str):
        m = re.match(r'^(\d+(?:\.\d+)*)(?:[.:])?\s+(.+)$', line)
        if not m:
            m = re.match(r'^(\d+)(?:[.:])?\s+(.+)$', line)
            if not m:
                if re.match(r'^(INTRODUCTION|OVERALL DESCRIPTION|SPECIFIC REQUIREMENTS|SYSTEM FEATURES)\.?$', line.strip(), flags=re.IGNORECASE):
                    return 1
                if re.match(
                    r'^(External Interface Requirements|Functional Requirements|Non-Functional Requirements)\.?$',
                    line.strip(),
                    flags=re.IGNORECASE,
                ):
                    return 2
                return None
        num = m.group(1)
        depth = 1 if num.endswith('.') else len(num.split('.'))
        return max(1, min(depth, 5))

    def _split_label_value(line: str):
        m = re.match(r'^([A-Za-z][A-Za-z0-9 /&()_-]{1,120})\s*:\s*(.*)$', line or "")
        if not m:
            return None, None
        return m.group(1).strip(), (m.group(2) or "").strip()

    def _add_page_number(footer_paragraph):
        footer_paragraph.alignment = 2  # right
        footer_paragraph.add_run("Page ")
        run = footer_paragraph.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = "PAGE"
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_end)
        footer_paragraph.add_run(" of ")
        run2 = footer_paragraph.add_run()
        fld_char_begin2 = OxmlElement('w:fldChar')
        fld_char_begin2.set(qn('w:fldCharType'), 'begin')
        instr_text2 = OxmlElement('w:instrText')
        instr_text2.set(qn('xml:space'), 'preserve')
        instr_text2.text = "NUMPAGES"
        fld_char_end2 = OxmlElement('w:fldChar')
        fld_char_end2.set(qn('w:fldCharType'), 'end')
        run2._r.append(fld_char_begin2)
        run2._r.append(instr_text2)
        run2._r.append(fld_char_end2)

    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)

    if doc.sections and doc.sections[0].footer.paragraphs:
        footer_p = doc.sections[0].footer.paragraphs[0]
    else:
        footer_p = doc.sections[0].footer.add_paragraph()
    _add_page_number(footer_p)

    display_title = (title or '').strip() or 'Untitled Project'
    author_label = (author or '').strip() or 'System'
    version_label = (version or '').strip() or '1.0'
    created_label = (date_value or '').strip() or datetime.now().strftime('%Y-%m-%d')
    doc_id_display = (document_id or '').strip() or '-'

    kicker = doc.add_paragraph('Software Requirements Specification')
    kicker.runs[0].font.size = Pt(9)
    kicker.runs[0].font.bold = True
    kicker.runs[0].font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    kicker.paragraph_format.space_after = Pt(4)

    h_title = doc.add_heading(display_title, level=1)
    h_title.runs[0].font.size = Pt(22)
    h_title.runs[0].font.bold = True

    sub = doc.add_paragraph(
        'Structured technical document aligned for formal review, sign-off, and handover.'
    )
    sub.runs[0].font.size = Pt(11)
    sub.runs[0].font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    sub.paragraph_format.space_after = Pt(10)

    meta_tbl = doc.add_table(rows=2, cols=4)
    meta_tbl.style = 'Table Grid'
    hdr = meta_tbl.rows[0].cells
    hdr[0].text = 'Document ID'
    hdr[1].text = 'Version'
    hdr[2].text = 'Date'
    hdr[3].text = 'Author'
    for c in hdr:
        for run in c.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    row = meta_tbl.rows[1].cells
    row[0].text = doc_id_display
    row[1].text = version_label
    row[2].text = created_label
    row[3].text = author_label
    doc.add_paragraph('')

    badges = doc.add_paragraph()
    badges.add_run('Req2Design').font.size = Pt(8)
    badges.add_run('   ')
    badges.add_run('IEEE 830-1998 aligned').font.size = Pt(8)
    badges.add_run('   ')
    badges.add_run('Generated draft for expert review').font.size = Pt(8)
    for r in badges.runs:
        r.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    doc.add_paragraph('')

    doc.add_page_break()

    source_text = str(raw_text or '').strip()
    if not source_text and isinstance(sections, dict):
        source_text = json.dumps(sections, indent=2)
    source_text = _normalize_for_export(source_text)

    lines = [ln.strip() for ln in source_text.splitlines() if ln.strip()]

    # Table of contents (manual list for stable cross-platform output)
    toc_entries = []
    for line in lines:
        depth = _heading_depth(line)
        if depth:
            toc_entries.append((depth, line))
    if toc_entries:
        doc.add_heading('Table of Contents', level=1)
        for depth, text_line in toc_entries:
            p = doc.add_paragraph(style='List Bullet')
            if depth > 1:
                p.paragraph_format.left_indent = None
            p.add_run(text_line)
        doc.add_page_break()

    seen_major = False
    for line in lines:
        depth = _heading_depth(line)
        if depth:
            # Put each major section heading on a new page.
            if depth == 1:
                if seen_major:
                    doc.add_page_break()
                seen_major = True
            # §3.2 Functional Requirements starts on its own page (matches PDF/HTML export).
            if depth == 2 and re.match(r'^3\.2[.:]?\s+Functional Requirements\b', line, flags=re.IGNORECASE):
                doc.add_page_break()
            level = 1 if depth == 1 else 2 if depth == 2 else 3
            doc.add_heading(line, level=level)
            continue

        frm = re.match(r'^(FR-\d+|NFR-\d+)\s*:?\s*(.*)$', line, flags=re.IGNORECASE)
        if frm:
            p = doc.add_paragraph()
            p.add_run(frm.group(1).upper()).bold = True
            if (frm.group(2) or '').strip():
                p.add_run(f" {frm.group(2).strip()}")
            continue

        if line.startswith(('- ', '* ')):
            bullet_text = line[2:].strip()
            k, v = _split_label_value(bullet_text)
            p = doc.add_paragraph(style='List Bullet')
            if k:
                p.add_run(f"{k}:").bold = True
                if v:
                    p.add_run(f" {v}")
            else:
                p.add_run(bullet_text)
            continue

        k, v = _split_label_value(line)
        if k:
            p = doc.add_paragraph()
            p.add_run(f"{k}:").bold = True
            if v:
                p.add_run(f" {v}")
            continue

        doc.add_paragraph(line)

    doc.add_paragraph('')
    end_hdr = doc.add_paragraph()
    end_run = end_hdr.add_run('End of document')
    end_run.bold = True
    end_run.font.size = Pt(10)
    end_run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    end_hdr.paragraph_format.space_after = Pt(4)
    end_note = doc.add_paragraph(
        'Generated by Req2Design – AI SRS Engineering Platform. This footer is added by the application '
        '(not the model) to keep exports consistent.'
    )
    end_note.runs[0].font.size = Pt(9)
    end_note.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    safe_title = re.sub(r'[\\/:*?"<>|]+', '', str(title or 'SRS')).strip() or 'SRS'
    safe_title = re.sub(r'\s+', '_', safe_title)
    out_path = os.path.join(tempfile.gettempdir(), f"{safe_title}.docx")
    doc.save(out_path)
    return out_path


@app.route('/api/generate-srs-docx', methods=['POST'])
def generate_srs_docx():
    """
    Generate SRS as DOCX and return as downloadable file.
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400

        out_path = _build_srs_docx(
            document_id=data.get('document_id', f"SRS-{datetime.now().strftime('%Y%m%d-%H%M%S')}"),
            title=data.get('title', 'Software Requirements Specification'),
            version=data.get('version', '1.0'),
            date_value=data.get('date', datetime.now().strftime('%Y-%m-%d')),
            author=data.get('author', 'System'),
            raw_text=data.get('raw_text', ''),
            sections=data.get('sections') or data.get('srs_sections') or {},
        )

        _resp = send_file(
            out_path,
            as_attachment=True,
            download_name=os.path.basename(out_path),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        _resp.headers['X-Export-Format'] = 'docx'
        return _resp
    except Exception as e:
        logger.error(f"Error generating SRS DOCX: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/generate-srs-from-audio', methods=['POST'])
def generate_srs_from_audio():
    """
    Direct SRS generation from audio without preprocessing/validation pipeline.
    Steps:
      1) Accept audio file
      2) Transcribe with Whisper
      3) Generate SRS via SRSModelGenerator using the transcription as input
    """
    temp_file_path = None
    try:
        audio_file = request.files.get('audio')
        project_info_str = request.form.get('project_info', '{}')
        try:
            project_info = json.loads(project_info_str) if project_info_str else {}
        except json.JSONDecodeError:
            project_info = {}

        if not audio_file:
            return jsonify({'error': 'No audio file provided'}), 400

        original_filename = audio_file.filename or 'recording'
        file_ext = 'webm'
        if '.' in original_filename:
            file_ext = original_filename.rsplit('.', 1)[1].lower()

        with tempfile.NamedTemporaryFile(suffix=f'.{file_ext}', delete=False) as temp_file:
            audio_file.save(temp_file.name)
            temp_file_path = temp_file.name

        # Verify file has content after saving
        if not os.path.exists(temp_file_path) or os.path.getsize(temp_file_path) == 0:
            return jsonify({'error': 'Audio file is empty. Please record again.'}), 400

        # Duration check
        try:
            duration = orchestrator.processor._get_audio_duration(temp_file_path)
            if duration <= 0.5:
                return jsonify({'error': 'Audio is too short. Please record at least 1 second.'}), 400
        except Exception as dur_err:
            logger.warning(f"Duration check failed: {dur_err}")

        # Transcribe
        if not getattr(orchestrator.processor, 'models_loaded', False):
            orchestrator.processor._load_models()
        transcription = orchestrator.processor._transcribe_audio(temp_file_path)

        if not transcription or not transcription.strip():
            return jsonify({'error': 'Transcription is empty. Please try again with clearer audio.'}), 400

        # Light deduplication: remove consecutive duplicate lines/paragraphs
        def _dedupe_text(text: str) -> str:
            lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
            deduped = []
            prev = None
            for ln in lines:
                if ln != prev:
                    deduped.append(ln)
                prev = ln
            # Also dedupe by paragraphs split on blank lines
            paragraphs = [p.strip() for p in "\n".join(deduped).split("\n\n") if p.strip()]
            deduped_paras = []
            prevp = None
            for p in paragraphs:
                if p != prevp:
                    deduped_paras.append(p)
                prevp = p
            return "\n\n".join(deduped_paras)

        transcription = _dedupe_text(transcription)

        # Generate SRS directly from transcription
        results = [{'original_text': transcription}]
        srs = _generate_srs_document(results, project_info)

        return jsonify(serialize_srs(srs))

    except Exception as e:
        logger.error(f"Error generating SRS from audio: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to generate SRS from audio: {str(e)}'}), 500
    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
            except Exception:
                pass


@app.route('/api/generate-srs-from-file', methods=['POST'])
def generate_srs_from_file():
    """
    Direct SRS generation from uploaded file (txt, docx or pdf).
    - Saves file locally
    - Extracts text (best effort for PDF)
    - Generates SRS without preprocessing pipeline
    """
    temp_file_path = None
    try:
        file = request.files.get('file')
        project_info_str = request.form.get('project_info', '{}')
        try:
            project_info = json.loads(project_info_str) if project_info_str else {}
        except json.JSONDecodeError:
            project_info = {}

        if not file:
            return jsonify({'error': 'No file provided'}), 400

        original_filename = file.filename or 'document'
        ext = original_filename.rsplit('.', 1)[-1].lower() if '.' in original_filename else 'txt'

        with tempfile.NamedTemporaryFile(suffix=f'.{ext}', delete=False) as temp_file:
            file.save(temp_file.name)
            temp_file_path = temp_file.name

        if not os.path.exists(temp_file_path) or os.path.getsize(temp_file_path) == 0:
            return jsonify({'error': 'File is empty. Please upload a valid file.'}), 400

        text_content = ""
        if ext == 'txt':
            with open(temp_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text_content = f.read()
        elif ext == 'pdf':
            text_content = _extract_text_from_pdf(temp_file_path)
        elif ext == 'docx':
            try:
                doc = Document(temp_file_path)
                text_content = "\n".join(p.text for p in doc.paragraphs if str(p.text).strip())
            except Exception:
                text_content = ""
        else:
            return jsonify({'error': 'Unsupported file type. Please upload .txt, .docx or .pdf'}), 400

        if not text_content or len(text_content.strip().split()) < 10:
            return jsonify({'error': 'File content is too short or could not be extracted. Please provide a valid text or PDF.'}), 400

        results = [{'original_text': text_content}]
        srs = _generate_srs_document(results, project_info)

        return jsonify(serialize_srs(srs))

    except Exception as e:
        logger.error(f"Error generating SRS from file: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to generate SRS from file: {str(e)}'}), 500
    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
            except Exception:
                pass

def _convert_raw_text_to_html(raw_text: str, document_id: str, title: str, version: str, date: str, author: str) -> str:
    """
    Convert raw SRS text from model to HTML, preserving exact formatting from model output.
    
    Args:
        raw_text: Raw text output from the SRS model
        document_id: Unique identifier for the SRS document
        title: Document title
        version: Document version
        date: Document date
        author: Document author
    
    Returns:
        HTML string with cleaned and formatted SRS content
    """
    # Remove markdown code blocks if present (but preserve the content)
    text = raw_text.strip()
    
    # Remove markdown code block markers (```plaintext, ```, etc.)
    if text.startswith("```"):
        # Remove opening code block markers (```plaintext, ```text, etc.)
        text = re.sub(r'^```[a-z]*\s*\n?', '', text, flags=re.IGNORECASE)
        # Remove closing code block markers
        text = re.sub(r'\n?```\s*$', '', text, flags=re.MULTILINE)
        text = text.strip()
    
    # Remove the disclaimer text at the bottom if present
    disclaimer_patterns = [
        r'This document adheres strictly to the IEEE 830-1998 format.*?specifications\.?\s*$',
        r'This document adheres.*?IEEE 830.*?specifications\.?\s*$',
        r'No additional content or assumptions have been added.*?specifications\.?\s*$',
    ]
    for pattern in disclaimer_patterns:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE | re.DOTALL | re.MULTILINE)
    
    # Keep only the standards body from the first canonical heading onward.
    # This removes any model preamble that appears before INTRODUCTION.
    intro_match = re.search(r'(?im)^\s*INTRODUCTION\s*$', text)
    if intro_match:
        text = text[intro_match.start():]

    # If the model accidentally appends a second SRS after "End of Document.",
    # keep only the first complete document block.
    end_marker = re.search(r'(?i)\bEnd of Document\.', text)
    if end_marker:
        text = text[:end_marker.end()]

    # Clean up any trailing whitespace or newlines
    text = text.strip()

    def _normalize_layout(t: str) -> str:
        y = t
        # Fix cases like "1Introduction" / "1.1Purpose" by inserting space after numeric prefix.
        y = re.sub(r'(?m)^(\d+(?:\.\d+)*)([A-Za-z])', r'\1 \2', y)
        # Rejoin hyphenated words split across line breaks (e.g., "Non-\nfunctional").
        y = re.sub(r'([A-Za-z])-\s*\n\s*([A-Za-z])', r'\1-\2', y)
        y = re.sub(r'(?m)^\s*-{3,}\s*$', '', y)
        y = re.sub(r'(?mi)^\s*===\s*[A-Z0-9 _/\-()]+(?:START|END)\s*===\s*$', '', y)
        y = re.sub(r'\s+(?=(INTRODUCTION|OVERALL DESCRIPTION|SPECIFIC REQUIREMENTS|SYSTEM FEATURES)\b)', '\n', y, flags=re.IGNORECASE)
        y = re.sub(r'\s+(?=\d+\.\s+[A-Z])', '\n', y)
        y = re.sub(r'\s+(?=\d+(?:\.\d+)+\s+[A-Z])', '\n', y)
        y = re.sub(r'\s+(?=FR-\d+\b)', '\n', y, flags=re.IGNORECASE)
        y = re.sub(r'\s+(?=NFR-\d+\b)', '\n', y, flags=re.IGNORECASE)
        labels = [
            "Purpose:",
            "Scope:",
            "Definitions/Acronyms:",
            "References:",
            "Overview:",
            "Product Perspective:",
            "Product Functions:",
            "User Characteristics:",
            "Constraints:",
            "Assumptions/Dependencies:",
            "External Interface Requirements:",
            "User Interface:",
            "Hardware Interface:",
            "Software Interface:",
            "Communication Interface:",
            "Functional Requirements:",
            "Non-functional Requirements:",
            "System Features:",
            "Input:",
            "Processing:",
            "Output:",
            "Priority:",
            "Feature 1:",
            "Feature 2:",
            "Feature 3:",
            "Feature 4:",
            "Feature 5:",
            "Feature 6:",
            "Feature 7:",
            "Feature 8:",
            "Feature 9:",
            "Feature 10:",
        ]
        for lbl in labels:
            y = re.sub(rf'\s+(?={re.escape(lbl)})', '\n', y, flags=re.IGNORECASE)
        y = re.sub(r'\s+-\s+', '\n- ', y)
        y = re.sub(r'(?<!\n)(Input|Processing|Output|Priority|Purpose|Scope|Definitions/Acronyms|References|Overview|Product Perspective|Product Functions|User Characteristics|Constraints|Assumptions/Dependencies|External Interface Requirements|Functional Requirements|Non-functional Requirements|System Features)\s*:', r'\n\1:', y, flags=re.IGNORECASE)
        y = re.sub(r'\n{3,}', '\n\n', y)
        # Remove stray single-dash lines introduced by model formatting.
        y = re.sub(r'(?m)^\s*-\s*$', '', y)
        # Remove duplicated ending lines and keep one canonical ending marker.
        y = re.sub(r'(?im)^(End of document\.)\s*$', r'\1', y)
        y = re.sub(r'(?im)^(End of document\.\s*){2,}$', 'End of document.', y)
        return _apply_ieee_numbering(y.strip())

    def _apply_ieee_numbering(text: str) -> str:
        lines = [ln.strip() for ln in text.splitlines()]
        out = []
        current_major = None
        current_sub = None

        major_map = {
            "introduction": (1, "Introduction"),
            "overall description": (2, "Overall Description"),
            "specific requirements": (3, "Specific Requirements"),
            "system features": (4, "System Features"),
        }
        sub_map = {
            1: {"purpose": (1, "Purpose"), "scope": (2, "Scope"), "definitions/acronyms": (3, "Definitions/Acronyms"), "references": (4, "References"), "overview": (5, "Overview")},
            2: {"product perspective": (1, "Product Perspective"), "product functions": (2, "Product Functions"), "user characteristics": (3, "User Characteristics"), "constraints": (4, "Constraints"), "assumptions/dependencies": (5, "Assumptions/Dependencies")},
            3: {"external interface requirements": (1, "External Interface Requirements"), "functional requirements": (2, "Functional Requirements"), "non-functional requirements": (3, "Non-Functional Requirements")},
        }
        subsub_map = {
            (3, 1): {"user interface": (1, "User Interface"), "hardware interface": (2, "Hardware Interface"), "software interface": (3, "Software Interface"), "communication interface": (4, "Communication Interface")},
            (3, 3): {"usability": (1, "Usability"), "reliability": (2, "Reliability"), "performance": (3, "Performance"), "portability": (4, "Portability")},
        }

        for line in lines:
            if not line:
                out.append("")
                continue
            if re.match(r'^\d+(?:\.\d+)*\s+', line):
                out.append(line)
                continue

            raw = line.strip().rstrip('.').rstrip(':')
            key = raw.lower()
            m = major_map.get(key)
            if m:
                current_major, title_txt = m
                current_sub = None
                out.append(f"{current_major}: {title_txt}")
                continue

            k = re.match(r'^([A-Za-z][A-Za-z0-9 /&()_-]{1,120})\s*:\s*(.*)$', line)
            if k and current_major in sub_map and k.group(1).strip().lower() in sub_map[current_major]:
                sub_idx, sub_label = sub_map[current_major][k.group(1).strip().lower()]
                current_sub = sub_idx
                tail = (k.group(2) or "").strip()
                if tail:
                    out.append(f"{current_major}.{sub_idx}: {sub_label}")
                    out.append(f"{sub_label}: {tail}")
                else:
                    out.append(f"{current_major}.{sub_idx}: {sub_label}")
                continue

            if k and (current_major, current_sub) in subsub_map and k.group(1).strip().lower() in subsub_map[(current_major, current_sub)]:
                subsub_idx, subsub_label = subsub_map[(current_major, current_sub)][k.group(1).strip().lower()]
                tail = (k.group(2) or "").strip()
                out.append(f"{current_major}.{current_sub}.{subsub_idx}: {subsub_label}" + (f": {tail}" if tail else ""))
                continue

            if current_major in sub_map and key in sub_map[current_major]:
                sub_idx, sub_label = sub_map[current_major][key]
                current_sub = sub_idx
                out.append(f"{current_major}.{sub_idx}: {sub_label}")
                continue
            if (current_major, current_sub) in subsub_map and key in subsub_map[(current_major, current_sub)]:
                subsub_idx, subsub_label = subsub_map[(current_major, current_sub)][key]
                out.append(f"{current_major}.{current_sub}.{subsub_idx}: {subsub_label}")
                continue

            out.append(line)

        return "\n".join(out).strip()

    text = _normalize_layout(text)

    # Remove common model separators + redundant header blocks (UI/PDF provides canonical header)
    text = re.sub(r'(?m)^\s*-{3,}\s*$', '', text)  # standalone --- lines
    text = re.sub(
        r'^\s*(?:Software Requirements Specification\s*\(SRS\)[^\n]*\n+)?(?:Author\s*:[^\n]*\n+)?(?:Date\s*:[^\n]*\n+)?\s*',
        '',
        text,
        flags=re.IGNORECASE,
    )
    text = text.strip()
    
    import html

    def _heading_depth(line: str):
        m = re.match(r'^(\d+(?:\.\d+)*)(?:[.:])?\s+(.+)$', line)
        if not m:
            m = re.match(r'^(\d+)(?:[.:])?\s+(.+)$', line)
            if not m:
                # Treat canonical major all-caps headings as top-level SRS headings.
                if re.match(r'^(INTRODUCTION|OVERALL DESCRIPTION|SPECIFIC REQUIREMENTS|SYSTEM FEATURES)\.?$', line.strip(), flags=re.IGNORECASE):
                    return 1
                if re.match(
                    r'^(External Interface Requirements|Functional Requirements|Non-Functional Requirements)\.?$',
                    line.strip(),
                    flags=re.IGNORECASE,
                ):
                    return 2
                return None
        num = m.group(1)
        depth = 1 if num.endswith('.') else len(num.split('.'))
        return max(1, min(depth, 5))

    def _format_srs_html_body(t: str):
        raw_rows = [r.strip() for r in t.split('\n')]
        # Merge accidental split lines like:
        # "Non-" + "functional Requirements:" -> "Non-functional Requirements:"
        rows = []
        i_merge = 0
        while i_merge < len(raw_rows):
            cur = raw_rows[i_merge]
            if (
                cur.lower() == "non-"
                and i_merge + 1 < len(raw_rows)
                and "functional requirements" in raw_rows[i_merge + 1].lower()
            ):
                rows.append("Non-Functional Requirements")
                i_merge += 2
                continue
            if (
                cur.endswith('-')
                and i_merge + 1 < len(raw_rows)
                and raw_rows[i_merge + 1]
                and re.match(r'^[a-z]', raw_rows[i_merge + 1])
            ):
                rows.append(f"{cur}{raw_rows[i_merge + 1]}")
                i_merge += 2
                continue
            rows.append(cur)
            i_merge += 1
        out = []
        toc = []
        in_list = False
        first_major_done = False
        heading_anchor_counts = {}

        def _slugify(value: str) -> str:
            base = re.sub(r'[^a-z0-9]+', '-', str(value).lower()).strip('-') or 'section'
            n = heading_anchor_counts.get(base, 0) + 1
            heading_anchor_counts[base] = n
            return f"{base}-{n}" if n > 1 else base

        def close_list():
            nonlocal in_list
            if in_list:
                out.append('</ul>')
                in_list = False

        i = 0
        n = len(rows)
        while i < n:
            row = rows[i]
            if not row:
                close_list()
                i += 1
                continue

            depth = _heading_depth(row)
            if depth:
                close_list()
                tag = 'h2' if depth == 1 else 'h3' if depth == 2 else 'h4'
                parsed = re.match(r'^(\d+(?:\.\d+)*)(?:[.:])?\s+(.+)$', row)
                if parsed:
                    num_txt = parsed.group(1)
                    label_txt = parsed.group(2).strip()
                    # If model merged heading and sentence (e.g., "PurposeThe ..."),
                    # split into clean heading + body text.
                    known_labels = [
                        "Purpose",
                        "Scope",
                        "Definitions/Acronyms",
                        "References",
                        "Overview",
                        "Product Perspective",
                        "Product Functions",
                        "User Characteristics",
                        "Constraints",
                        "Assumptions/Dependencies",
                        "External Interface Requirements",
                        "Functional Requirements",
                        "Non-Functional Requirements",
                        "System Features",
                        "Usability",
                        "Reliability",
                        "Performance",
                        "Portability",
                    ]
                    merged_tail = ""
                    for lbl in known_labels:
                        if label_txt.startswith(lbl):
                            rest = label_txt[len(lbl):].strip()
                            if rest:
                                merged_tail = rest
                                label_txt = lbl
                            break
                    display_txt = f"{num_txt}. {label_txt}"
                    anchor = _slugify(display_txt)
                    num = html.escape(num_txt)
                    label = html.escape(label_txt)
                    toc_label = re.sub(r':\s*.+$', '', display_txt).strip()
                    toc.append((depth, toc_label, anchor))
                    heading_cls = f'srs-h d{depth}'
                    if depth == 1:
                        if first_major_done:
                            heading_cls += ' section-break'
                        first_major_done = True
                    if re.match(r'^Non-Functional Requirements\b', display_txt.split('. ', 1)[-1], flags=re.IGNORECASE):
                        heading_cls += ' section-break'
                    if num_txt == '3.2' and re.match(r'^Functional Requirements\b', label_txt, flags=re.IGNORECASE):
                        heading_cls += ' section-break'
                    out.append(
                        f'<a name="{anchor}"></a><{tag} id="{anchor}" class="{heading_cls}">{num}: {label}</{tag}>'
                    )
                    if merged_tail:
                        out.append(
                            f'<p class="srs-p kv"><span class="k">{html.escape(label_txt)}:</span> '
                            f'<span>{html.escape(merged_tail)}</span></p>'
                        )
                else:
                    display_txt = row
                    anchor = _slugify(display_txt)
                    toc_label = re.sub(r':\s*.+$', '', display_txt).strip()
                    toc.append((depth, toc_label, anchor))
                    heading_cls = f'srs-h d{depth}'
                    if depth == 1:
                        if first_major_done:
                            heading_cls += ' section-break'
                        first_major_done = True
                    if re.match(r'^Non-Functional Requirements\.?$', display_txt, flags=re.IGNORECASE):
                        heading_cls += ' section-break'
                    if re.match(r'^Functional Requirements\.?$', display_txt, flags=re.IGNORECASE):
                        heading_cls += ' section-break'
                    if re.match(r'^3\.2[.:]?\s+Functional Requirements\b', display_txt, flags=re.IGNORECASE):
                        heading_cls += ' section-break'
                    out.append(f'<a name="{anchor}"></a><{tag} id="{anchor}" class="{heading_cls}">{html.escape(row)}</{tag}>')

                # Structured FR table
                if re.match(r'^Functional Requirements$', row, flags=re.IGNORECASE):
                    fr_rows = []
                    j = i + 1
                    while j < n:
                        line = rows[j]
                        if not line:
                            j += 1
                            continue
                        if _heading_depth(line):
                            break
                        mfr = re.match(r'^(FR-\d+)\s*:?\s*(.*)$', line, flags=re.IGNORECASE)
                        if mfr:
                            req_id = mfr.group(1).upper()
                            title = (mfr.group(2) or '').strip()
                            bullets = []
                            priority = ""
                            j += 1
                            while j < n:
                                inner = rows[j]
                                if not inner:
                                    j += 1
                                    continue
                                if _heading_depth(inner) or re.match(r'^(FR-\d+)\s*:?\s*', inner, flags=re.IGNORECASE):
                                    break
                                pri = re.match(r'^-?\s*Priority\s*:\s*(.+)$', inner, flags=re.IGNORECASE)
                                if pri:
                                    priority = pri.group(1).strip()
                                else:
                                    bl = re.sub(r'^[-*]\s*', '', inner).strip()
                                    if bl:
                                        bullets.append(bl)
                                j += 1
                            description = title or (bullets[0] if bullets else '')
                            acceptance = bullets[0] if bullets else (title or '')
                            rationale = "Supports required business capability and user workflow."
                            source = "Stakeholder requirements input"
                            fr_rows.append(
                                (req_id, description, rationale, source, acceptance, priority or "Medium")
                            )
                            continue
                        j += 1
                    if fr_rows:
                        out.append(
                            '<table class="req-table"><thead><tr>'
                            '<th style="width:10%;">ID</th><th style="width:25%;">Description</th>'
                            '<th style="width:20%;">Rationale</th><th style="width:15%;">Source</th>'
                            '<th style="width:20%;">Acceptance Criteria</th><th style="width:10%;">Priority</th>'
                            '</tr></thead><tbody>'
                        )
                        for rid, desc, rat, src, acc, pri in fr_rows:
                            out.append(
                                f'<tr><td>{html.escape(rid)}</td><td>{html.escape(desc)}</td>'
                                f'<td>{html.escape(rat)}</td><td>{html.escape(src)}</td>'
                                f'<td>{html.escape(acc)}</td><td>{html.escape(pri)}</td></tr>'
                            )
                        out.append('</tbody></table>')
                        i = j
                        continue

                # Structured NFR table
                if re.match(r'^Non-Functional Requirements$', row, flags=re.IGNORECASE):
                    nfr_rows = []
                    j = i + 1
                    category = None
                    points = []
                    nfr_id = 1

                    def _flush_nfr(cat, items):
                        nonlocal nfr_id
                        if not cat:
                            return
                        desc = '; '.join(items).strip()
                        if not desc:
                            return
                        nfr_rows.append(
                            (
                                f"NFR-{nfr_id:02d}",
                                f"{cat}: {desc}",
                                f"Ensures {cat.lower()} quality objectives are met.",
                                "Quality attributes baseline",
                                items[0] if items else desc,
                                "High",
                            )
                        )
                        nfr_id += 1

                    while j < n:
                        line = rows[j]
                        if not line:
                            j += 1
                            continue
                        if _heading_depth(line):
                            break
                        cat_match = re.match(r'^([A-Za-z][A-Za-z0-9 /&()_-]{1,80})\s*:\s*$', line)
                        if cat_match:
                            _flush_nfr(category, points)
                            category = cat_match.group(1).strip()
                            points = []
                            j += 1
                            continue
                        bl = re.sub(r'^[-*]\s*', '', line).strip()
                        if bl:
                            points.append(bl)
                        j += 1
                    _flush_nfr(category, points)
                    if nfr_rows:
                        out.append(
                            '<table class="req-table"><thead><tr>'
                            '<th style="width:10%;">ID</th><th style="width:25%;">Description</th>'
                            '<th style="width:20%;">Rationale</th><th style="width:15%;">Source</th>'
                            '<th style="width:20%;">Acceptance Criteria</th><th style="width:10%;">Priority</th>'
                            '</tr></thead><tbody>'
                        )
                        for rid, desc, rat, src, acc, pri in nfr_rows:
                            out.append(
                                f'<tr><td>{html.escape(rid)}</td><td>{html.escape(desc)}</td>'
                                f'<td>{html.escape(rat)}</td><td>{html.escape(src)}</td>'
                                f'<td>{html.escape(acc)}</td><td>{html.escape(pri)}</td></tr>'
                            )
                        out.append('</tbody></table>')
                        i = j
                        continue

                i += 1
                continue

            frm = re.match(r'^(FR-\d+|NFR-\d+)\s*:?\s*(.*)$', row, flags=re.IGNORECASE)
            if frm:
                close_list()
                out.append(
                    f'<p class="srs-p fr"><span class="id">{html.escape(frm.group(1).upper())}</span> '
                    f'<span>{html.escape(frm.group(2))}</span></p>'
                )
                i += 1
                continue

            kv = re.match(r'^([A-Za-z][A-Za-z0-9 /&()_-]{1,120})\s*:\s*(.*)$', row)
            if kv:
                close_list()
                raw_key = kv.group(1).strip()
                raw_value = (kv.group(2) or '').strip()

                # Promote section-style label lines to real headings.
                if re.match(r'^non[- ]functional requirements$', raw_key, flags=re.IGNORECASE):
                    heading_text = "Non-Functional Requirements"
                    anchor = _slugify(heading_text)
                    toc.append((2, heading_text, anchor))
                    out.append(f'<a name="{anchor}"></a><h3 id="{anchor}" class="srs-h d2 section-break">{heading_text}</h3>')
                    i += 1
                    continue
                if re.match(r'^functional requirements$', raw_key, flags=re.IGNORECASE):
                    heading_text = "Functional Requirements"
                    anchor = _slugify(heading_text)
                    toc.append((2, heading_text, anchor))
                    out.append(f'<a name="{anchor}"></a><h3 id="{anchor}" class="srs-h d2 section-break">{heading_text}</h3>')
                    i += 1
                    continue
                if re.match(r'^external interface requirements$', raw_key, flags=re.IGNORECASE):
                    heading_text = "External Interface Requirements"
                    anchor = _slugify(heading_text)
                    toc.append((2, heading_text, anchor))
                    out.append(f'<a name="{anchor}"></a><h3 id="{anchor}" class="srs-h d2">{heading_text}</h3>')
                    i += 1
                    continue

                key = html.escape(raw_key)
                value = html.escape(raw_value)
                if value.strip():
                    out.append(
                        f'<p class="srs-p kv"><span class="k">{key}:</span> '
                        f'<span>{value}</span></p>'
                    )
                else:
                    out.append(f'<p class="srs-p kv"><span class="k">{key}:</span></p>')
                i += 1
                continue

            if row.startswith('- ') or row.startswith('* '):
                if not in_list:
                    out.append('<ul class="srs-list">')
                    in_list = True
                bullet_text = row[2:].strip()
                bkv = re.match(r'^([A-Za-z][A-Za-z0-9 /&()_-]{1,80})\s*:\s*(.*)$', bullet_text)
                if bkv:
                    key = html.escape(bkv.group(1))
                    value = html.escape(bkv.group(2) or '')
                    if value.strip():
                        out.append(f'<li><span class="k">{key}:</span> <span>{value}</span></li>')
                    else:
                        out.append(f'<li><span class="k">{key}:</span></li>')
                else:
                    out.append(f'<li>{html.escape(bullet_text)}</li>')
                i += 1
                continue

            close_list()
            out.append(f'<p class="srs-p">{html.escape(row)}</p>')
            i += 1

        close_list()
        return ''.join(out), toc

    body_html, toc_entries = _format_srs_html_body(text)
    toc_items = [
        (1, "Table of Contents", "table-of-contents"),
        (1, "Revision History", "revision-history"),
        *toc_entries,
    ]
    toc_html = ''.join(
        (
            f'<li class="d{depth}">'
            f'<a href="#{html.escape(anchor)}">'
            f'<span class="toc-label">{html.escape(label)}</span>'
            f'<span class="toc-dots" aria-hidden="true"></span>'
            f'</a>'
            f'</li>'
        )
        for depth, label, anchor in toc_items
    )

    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    author_label = (author or '').strip() or 'System'
    version_label = (version or '').strip() or '1.0'
    created_label = (date or '').strip() or datetime.now().strftime("%Y-%m-%d")
    return f"""<!DOCTYPE html>
<html>
<head>
    <title>{html.escape(title)}</title>
    <meta charset="UTF-8">
    <style>
        @page {{
            size: A4;
            margin: 24mm 18mm 20mm 18mm;
            @bottom-right {{
                content: "Page " counter(page) " of " counter(pages);
                font-size: 9pt;
                color: #475569;
            }}
        }}
        body {{
            font-family: 'Times New Roman', Times, serif;
            margin: 0;
            line-height: 1.7;
            color: #0f172a;
            font-size: 12pt;
        }}
        .doc-header {{
            border: 1px solid #cbd5e1;
            border-radius: 8px;
            overflow: hidden;
            margin-bottom: 16px;
        }}
        /* Title page — matches Req2Design SRS viewer export styling */
        .pdf-title-cover {{
            border: 1px solid #d7e0ee;
            border-radius: 14px;
            padding: 12mm 10mm;
            margin: 0 0 12mm;
            page-break-after: always;
            background: linear-gradient(135deg, rgba(99, 102, 241, 0.08), rgba(15, 23, 42, 0.03)), #f8fbff;
            box-sizing: border-box;
        }}
        .srs-cover-kicker {{
            letter-spacing: 0.14em;
            text-transform: uppercase;
            font-size: 9pt;
            font-weight: 700;
            color: #475569;
            margin: 0 0 6pt;
        }}
        .srs-cover-title {{
            font-size: 22pt;
            font-weight: 800;
            margin: 0 0 8pt;
            color: #0f172a;
            line-height: 1.15;
        }}
        .srs-cover-sub {{
            font-size: 11pt;
            color: #334155;
            margin: 0 0 12pt;
            line-height: 1.45;
        }}
        .srs-cover-meta-table {{
            width: 100%;
            border-collapse: collapse;
            font-size: 10pt;
            margin-top: 6pt;
        }}
        .srs-cover-meta-table th,
        .srs-cover-meta-table td {{
            border: 1px solid #d7e0ee;
            padding: 7pt 8pt;
            text-align: left;
            vertical-align: top;
        }}
        .srs-cover-meta-table th {{
            font-size: 8pt;
            text-transform: uppercase;
            letter-spacing: 0.08em;
            color: #64748b;
            background: rgba(255, 255, 255, 0.9);
            font-weight: 700;
        }}
        .srs-cover-meta-table td {{
            font-weight: 600;
            color: #0f172a;
        }}
        .badge-row {{
            margin-top: 10pt;
            font-size: 8pt;
            color: #475569;
        }}
        .badge {{
            display: inline-block;
            border: 1px solid #cbd5e1;
            border-radius: 999px;
            padding: 4pt 8pt;
            margin: 4pt 6pt 0 0;
            background: rgba(255, 255, 255, 0.88);
        }}
        .pdf-title-foot {{
            margin-top: 12pt;
            font-size: 9pt;
            color: #64748b;
            font-style: italic;
        }}
        .doc-header-top {{
            padding: 13px 15px;
            border-bottom: 1px solid #cbd5e1;
        }}
        .doc-title {{
            margin: 0;
            font-size: 24pt;
            font-weight: 800;
            color: #0f172a;
        }}
        .doc-meta {{
            margin-top: 8px;
            font-size: 10.5pt;
            color: #475569;
            display: grid;
            grid-template-columns: repeat(2, minmax(0, 1fr));
            gap: 4px 12px;
        }}
        .doc-strip {{
            padding: 8px 15px;
            background: #f1f5f9;
            font-size: 10.5pt;
            color: #64748b;
        }}
        .srs-h {{
            margin: 16px 0 8px;
            line-height: 1.3;
            font-weight: 800;
            page-break-after: avoid;
            display: flex;
            align-items: baseline;
            gap: 6px;
        }}
        .srs-h.d1 {{
            font-size: 20pt;
            border-bottom: 1px solid #dbe3ee;
            padding-bottom: 4px;
        }}
        .srs-h.section-break {{ page-break-before: always; }}
        .srs-h.d2 {{ font-size: 16pt; }}
        .srs-h.d3, .srs-h.d4, .srs-h.d5 {{ font-size: 14pt; }}
        .srs-h .h-num {{
            min-width: 34pt;
            display: inline-block;
            font-variant-numeric: tabular-nums;
            font-weight: 700;
        }}
        .srs-h.d2 .h-num {{ min-width: 44pt; }}
        .srs-h.d3 .h-num, .srs-h.d4 .h-num, .srs-h.d5 .h-num {{ min-width: 52pt; }}
        .srs-p {{
            margin: 4px 0 10px;
            text-align: left;
        }}
        .srs-p.kv .k, .srs-p.fr .id {{ font-weight: 700; }}
        .srs-list {{
            margin: 4px 0 10px 20px;
            padding-left: 4px;
        }}
        .srs-list li {{ margin: 4px 0; }}
        .srs-doc-root {{
            font-size: 12pt;
            line-height: 1.92;
        }}
        .req-table {{
            width: 100%;
            border-collapse: collapse;
            font-size: 10.2pt;
            margin: 8px 0 14px;
        }}
        .req-table th, .req-table td {{
            border: 1px solid #d7e0ee;
            padding: 6px 7px;
            text-align: left;
            vertical-align: top;
        }}
        .req-table th {{
            background: rgba(248, 250, 252, 0.95);
            font-weight: 700;
            color: #334155;
        }}
        .doc-footer {{
            margin-top: 18px;
            border: 1px solid #e2e8f0;
            border-radius: 10px;
            padding: 12px 14px;
            background: rgba(255, 255, 255, 0.96);
            color: #64748b;
            font-size: 9.5pt;
            line-height: 1.45;
        }}
        .doc-footer-title {{
            margin: 0 0 6px;
            font-size: 10.5pt;
            color: #0f172a;
        }}
        .doc-footer-note {{
            margin: 0;
        }}
        .doc-footer-ts {{
            margin: 8px 0 0;
            font-size: 9pt;
            color: #94a3b8;
        }}
        .page-number {{
            position: fixed;
            right: 0;
            bottom: -2mm;
            font-size: 9pt;
            color: #64748b;
        }}
        .toc-box {{
            border: 1px solid #cbd5e1;
            border-radius: 8px;
            padding: 12px 14px;
            margin: 0 0 16px;
            page-break-before: always;
            page-break-after: always;
            min-height: 235mm;
            box-sizing: border-box;
        }}
        .toc-title {{
            margin: 0 0 8px;
            font-size: 14pt;
            font-weight: 700;
        }}
        .toc-list {{
            margin: 0;
            padding: 0;
            line-height: 1.45;
            list-style: none;
        }}
        .toc-list li {{ margin: 1px 0; }}
        .toc-list li a {{
            color: #000;
            text-decoration: none;
            display: flex;
            align-items: baseline;
            gap: 6px;
        }}
        .toc-list li a:hover {{
            text-decoration: underline;
        }}
        .toc-list li .toc-label {{
            white-space: nowrap;
        }}
        .toc-list li .toc-dots {{
            flex: 1;
            border-bottom: 1px dotted #111827;
            margin-bottom: 2px;
            min-width: 20px;
        }}
        .toc-list li a::after {{
            content: target-counter(attr(href), page);
            white-space: nowrap;
            min-width: 16px;
            text-align: right;
        }}
        .toc-list li.d2 {{ margin-left: 10px; }}
        .toc-list li.d3, .toc-list li.d4, .toc-list li.d5 {{ margin-left: 20px; }}
        .content {{ page-break-before: always; }}
        .revision-box {{
            margin: 0 0 18px;
            page-break-after: always;
        }}
        .revision-title {{
            margin: 0 0 8px;
            font-size: 16pt;
            font-weight: 700;
        }}
        .revision-table {{
            width: 100%;
            border-collapse: collapse;
            font-size: 11pt;
        }}
        .revision-table th, .revision-table td {{
            border: 1px solid #111827;
            padding: 6px;
            text-align: left;
            vertical-align: top;
        }}
    </style>
</head>
<body>
    <section class="pdf-title-cover">
      <p class="srs-cover-kicker">Software Requirements Specification</p>
      <h1 class="srs-cover-title">{html.escape((title or '').strip() or 'Untitled Project')}</h1>
      <p class="srs-cover-sub">Structured technical document aligned for formal review, sign-off, and handover.</p>
      <table class="srs-cover-meta-table" role="presentation">
        <thead>
          <tr>
            <th style="width: 28%;">Document ID</th>
            <th style="width: 18%;">Version</th>
            <th style="width: 22%;">Date</th>
            <th style="width: 32%;">Author</th>
          </tr>
        </thead>
        <tbody>
          <tr>
            <td>{html.escape((document_id or '').strip() or '-')}</td>
            <td>{html.escape(version_label)}</td>
            <td>{html.escape(created_label)}</td>
            <td>{html.escape(author_label)}</td>
          </tr>
        </tbody>
      </table>
      <div class="badge-row">
        <span class="badge">Req2Design</span>
        <span class="badge">IEEE 830-1998 aligned</span>
        <span class="badge">Generated draft for expert review</span>
      </div>
    </section>
    <pdf:nextpage />

    <section class="toc-box" id="table-of-contents">
      <h2 class="toc-title">Table of Contents</h2>
      <ul class="toc-list">
        {toc_html}
      </ul>
    </section>
    <pdf:nextpage />

    <section class="revision-box" id="revision-history">
      <h2 class="revision-title">Revision History</h2>
      <table class="revision-table" role="presentation">
        <thead>
          <tr>
            <th style="width: 22%;">Name</th>
            <th style="width: 18%;">Date</th>
            <th>Reason For Changes</th>
            <th style="width: 17%;">Version</th>
          </tr>
        </thead>
        <tbody>
          <tr>
            <td>{html.escape(author_label)}</td>
            <td>{html.escape(created_label)}</td>
            <td>Initial generated draft for review.</td>
            <td>{html.escape(version_label)}</td>
          </tr>
          <tr><td>&nbsp;</td><td>&nbsp;</td><td>&nbsp;</td><td>&nbsp;</td></tr>
          <tr><td>&nbsp;</td><td>&nbsp;</td><td>&nbsp;</td><td>&nbsp;</td></tr>
        </tbody>
      </table>
    </section>
    <pdf:nextpage />

    <div class="content srs-paper-shell">
      <div class="srs-paper">
        <div class="srs-doc-root">
          {body_html}
        </div>
        <div class="doc-footer">
          <p class="doc-footer-title"><strong>End of document</strong></p>
          <p class="doc-footer-note">Generated by <strong>Req2Design – AI SRS Engineering Platform</strong>. This footer is added by the application (not the model) to keep exports consistent.</p>
          <p class="doc-footer-ts">{html.escape(generated_at)}</p>
        </div>
      </div>
    </div>
    <div class="page-number">Page <pdf:pagenumber /> of <pdf:pagecount /></div>
</body>
</html>"""

def generate_html_content(srs_data: dict) -> str:
    """
    Generate HTML content for SRS document from structured data.
    
    Converts SRS document sections into a formatted HTML document suitable
    for display or printing. Handles both 'sections' and 'srs_sections' keys
    for backward compatibility.
    
    Args:
        srs_data: Dictionary containing SRS document data with:
            - title: Document title
            - document_id: Unique document identifier
            - version: Document version
            - date: Document date
            - author: Document author
            - sections or srs_sections: Dictionary containing parsed SRS sections
    
    Returns:
        HTML string containing the formatted SRS document
    """
    sections = srs_data.get('sections') or srs_data.get('srs_sections') or {}
    intro = sections.get('introduction', {})
    overall = sections.get('overall_description', {})
    return f"""
<!DOCTYPE html>
<html>
<head>
    <title>{srs_data['title']}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
        h2 {{ color: #34495e; margin-top: 30px; }}
        h3 {{ color: #7f8c8d; margin-top: 20px; }}
        .metadata {{ background-color: #ecf0f1; padding: 15px; margin-bottom: 20px; border-radius: 5px; }}
        ul {{ margin: 10px 0; }}
        li {{ margin: 5px 0; }}
        .section {{ margin: 20px 0; }}
    </style>
</head>
<body>
    <h1>{srs_data['title']}</h1>
    
    <div class="metadata">
        <p><strong>Document ID:</strong> {srs_data['document_id']}</p>
        <p><strong>Version:</strong> {srs_data['version']}</p>
        <p><strong>Date:</strong> {srs_data['date']}</p>
        <p><strong>Author:</strong> {srs_data['author']}</p>
    </div>
    
    <div class="section">
        <h2>1. Introduction</h2>
        <h3>1.1 Purpose</h3>
        <p>{intro.get('purpose','')}</p>
        
        <h3>1.2 Scope</h3>
        <p>{intro.get('scope','')}</p>
        
        <h3>1.3 Definitions</h3>
        <ul>
            {''.join(f'<li>{defn}</li>' for defn in intro.get('definitions', []))}
        </ul>
        
        <h3>1.4 Overview</h3>
        <p>{intro.get('overview','')}</p>
    </div>
    
    <div class="section">
        <h2>2. Overall Description</h2>
        <h3>2.1 Product Functions</h3>
        <ul>
            {''.join(f'<li>{func}</li>' for func in overall.get('product_functions', []))}
        </ul>
        
        <h3>2.2 User Characteristics</h3>
        <ul>
            {''.join(f'<li>{user}</li>' for user in overall.get('user_characteristics', []))}
        </ul>
        
        <h3>2.3 Constraints</h3>
        <ul>
            {''.join(f'<li>{constraint}</li>' for constraint in overall.get('constraints', []))}
        </ul>
        
        <h3>2.4 Assumptions</h3>
        <ul>
            {''.join(f'<li>{assumption}</li>' for assumption in (overall.get('assumptions', []) if isinstance(overall.get('assumptions'), list) else [overall.get('assumptions')] if overall.get('assumptions') else []))}
        </ul>
        
        <h3>2.5 Dependencies</h3>
        <ul>
            {''.join(f'<li>{dep}</li>' for dep in overall.get('dependencies', []))}
        </ul>
    </div>
    
    <div class="section">
        <h2>3. Note</h2>
        <p><em>This is an initial SRS document generated by Module 1. It contains only the Introduction and Overall Description sections. 
        Specific Requirements and other detailed sections will be generated in subsequent modules of the requirements engineering system.</em></p>
    </div>
</body>
</html>"""

@app.route('/api/expert-review/submit', methods=['POST'])
def expert_review_submit():
    """
    Submit an SRS snapshot for human expert review.
    Body: submitter (optional dict: user_id, username, email), srs_snapshot (object), requester_notes (optional str).
    """
    try:
        data = request.get_json(silent=True)
        if data is None or not isinstance(data, dict):
            return jsonify({"error": "Expected a JSON object body"}), 400
        srs_snapshot = data.get("srs_snapshot")
        if not isinstance(srs_snapshot, dict):
            return jsonify({"error": "srs_snapshot must be an object"}), 400
        doc_id = str(srs_snapshot.get("document_id") or srs_snapshot.get("id") or "").strip()
        if not doc_id:
            return jsonify({"error": "srs_snapshot must include document_id"}), 400

        submitter = data.get("submitter") if isinstance(data.get("submitter"), dict) else {}
        notes = sanitize_user_input(str(data.get("requester_notes") or "").strip())[:8000]

        raw_text = srs_snapshot.get("raw_text")
        if isinstance(raw_text, str) and len(raw_text) > 200_000:
            srs_snapshot = {**srs_snapshot, "raw_text": raw_text[:200_000] + "\n\n[truncated for review payload]"}

        rid = f"er_{uuid.uuid4().hex[:16]}"
        entry = {
            "id": rid,
            "status": "pending",
            "submitted_at": datetime.utcnow().isoformat() + "Z",
            "submitter": {
                "user_id": str(submitter.get("user_id") or "") or None,
                "username": str(submitter.get("username") or "") or None,
                "email": str(submitter.get("email") or "") or None,
            },
            "requester_notes": notes or None,
            "srs_snapshot": srs_snapshot,
            "review": None,
            "chat_messages": [],
        }

        with _expert_reviews_lock:
            reviews = _load_expert_reviews()
            reviews.insert(0, entry)
            if not _save_expert_reviews(reviews):
                return jsonify({"error": "Could not save review queue. Check server disk space and permissions."}), 500

        return jsonify({"ok": True, "id": rid, "status": "pending"})
    except Exception as e:
        logger.error("expert_review_submit: %s", e)
        return jsonify({"error": str(e)}), 500


@app.route('/api/expert-review/requests', methods=['GET'])
def expert_review_list():
    """List review requests. Query: status=pending|reviewed|all (default all)."""
    try:
        status_filter = str(request.args.get("status") or "all").strip().lower()
        if status_filter not in {"pending", "reviewed", "all"}:
            return jsonify({"error": "Invalid status; use pending, reviewed, or all"}), 400
        with _expert_reviews_lock:
            reviews = list(_load_expert_reviews())
        if status_filter in {"pending", "reviewed"}:
            reviews = [r for r in reviews if isinstance(r, dict) and r.get("status") == status_filter]
        reviews = [_normalize_review_entry(r) for r in reviews if isinstance(r, dict)]
        return jsonify({"requests": reviews})
    except Exception as e:
        logger.error("expert_review_list: %s", e)
        return jsonify({"error": str(e)}), 500


@app.route('/api/expert-review/requests/<rid>', methods=['GET'])
def expert_review_get(rid):
    try:
        rid = str(rid or "").strip()
        if not rid or len(rid) > 64 or not re.match(r"^er_[a-f0-9]+$", rid, re.I):
            return jsonify({"error": "Invalid review id"}), 400
        with _expert_reviews_lock:
            reviews = _load_expert_reviews()
        idx = _find_review_by_id(reviews, rid)
        if idx is None:
            return jsonify({"error": "Not found"}), 404
        return jsonify(_normalize_review_entry(reviews[idx]))
    except Exception as e:
        logger.error("expert_review_get: %s", e)
        return jsonify({"error": str(e)}), 500


@app.route('/api/expert-review/requests/<rid>/messages', methods=['POST'])
def expert_review_post_message(rid):
    """
    Append a chat message on a review thread. Body: sender_role (user|expert), body (str),
    author_label (optional), submitter (optional dict with user_id for user messages).
    """
    try:
        rid = str(rid or "").strip()
        if not rid or len(rid) > 64 or not re.match(r"^er_[a-f0-9]+$", rid, re.I):
            return jsonify({"error": "Invalid review id"}), 400
        data = request.get_json(silent=True)
        if data is None or not isinstance(data, dict):
            return jsonify({"error": "Expected a JSON object body"}), 400
        sender_role = str(data.get("sender_role") or "").strip().lower()
        if sender_role not in {"user", "expert"}:
            return jsonify({"error": "sender_role must be user or expert"}), 400
        body = sanitize_user_input(str(data.get("body") or "").strip(), max_length=8000)
        if len(body) < 1:
            return jsonify({"error": "body is required"}), 400
        author_label = str(data.get("author_label") or "").strip()[:200] or None
        req_submitter = data.get("submitter") if isinstance(data.get("submitter"), dict) else {}

        with _expert_reviews_lock:
            reviews = _load_expert_reviews()
            idx = _find_review_by_id(reviews, rid)
            if idx is None:
                return jsonify({"error": "Not found"}), 404
            item = reviews[idx]
            if not isinstance(item, dict):
                return jsonify({"error": "Invalid record"}), 500
            item = _normalize_review_entry(dict(item))
            stored_uid = (item.get("submitter") or {}).get("user_id")
            if sender_role == "user" and stored_uid:
                client_uid = str(req_submitter.get("user_id") or "").strip()
                if client_uid != str(stored_uid).strip():
                    return jsonify({"error": "User identity does not match this submission"}), 403

            mid = f"cm_{uuid.uuid4().hex[:16]}"
            msg = {
                "id": mid,
                "sent_at": datetime.utcnow().isoformat() + "Z",
                "sender_role": sender_role,
                "body": body,
            }
            if author_label:
                msg["author_label"] = author_label

            cm = list(item.get("chat_messages") or [])
            cm.append(msg)
            item["chat_messages"] = cm
            reviews[idx] = item
            if not _save_expert_reviews(reviews):
                return jsonify({"error": "Could not save message. Check server disk space and permissions."}), 500

        return jsonify({"ok": True, "message": msg})
    except Exception as e:
        logger.error("expert_review_post_message: %s", e)
        return jsonify({"error": str(e)}), 500


@app.route('/api/expert-review/requests/<rid>', methods=['PATCH'])
def expert_review_complete(rid):
    """
    Expert completes a review. Body: expert_feedback (str), verdict (optional: approved|needs_revision|rejected),
    expert_name (optional). In production, restrict to authenticated expert roles.
    """
    try:
        rid = str(rid or "").strip()
        if not rid or len(rid) > 64 or not re.match(r"^er_[a-f0-9]+$", rid, re.I):
            return jsonify({"error": "Invalid review id"}), 400
        data = request.get_json(silent=True)
        if data is None or not isinstance(data, dict):
            return jsonify({"error": "Expected a JSON object body"}), 400
        feedback = str(data.get("expert_feedback") or "").strip()
        if len(feedback) < 3:
            return jsonify({"error": "expert_feedback is required (at least a few characters)"}), 400
        if len(feedback) > 16000:
            feedback = feedback[:16000]
        verdict = str(data.get("verdict") or "approved").strip().lower()
        if verdict not in {"approved", "needs_revision", "rejected"}:
            verdict = "approved"
        expert_name = str(data.get("expert_name") or "").strip() or "Expert"

        with _expert_reviews_lock:
            reviews = _load_expert_reviews()
            idx = _find_review_by_id(reviews, rid)
            if idx is None:
                return jsonify({"error": "Not found"}), 404
            item = _normalize_review_entry(dict(reviews[idx]))
            if not isinstance(item, dict):
                return jsonify({"error": "Invalid record"}), 500
            if item.get("status") != "pending":
                return jsonify({"error": "This request is already reviewed"}), 409
            item["status"] = "reviewed"
            item["review"] = {
                "expert_feedback": feedback,
                "verdict": verdict,
                "expert_name": expert_name,
                "reviewed_at": datetime.utcnow().isoformat() + "Z",
            }
            reviews[idx] = item
            if not _save_expert_reviews(reviews):
                return jsonify({"error": "Could not save review. Check server disk space and permissions."}), 500

        return jsonify({"ok": True, "id": rid, "status": "reviewed"})
    except Exception as e:
        logger.error("expert_review_complete: %s", e)
        return jsonify({"error": str(e)}), 500


@app.route('/api/stats', methods=['GET'])
def get_system_stats():
    """Get system statistics"""
    try:
        stats = orchestrator.get_system_stats()
        return jsonify(stats)
    except Exception as e:
        logger.error(f"Error getting system stats: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/cleanup', methods=['POST'])
def cleanup_system():
    """Clean up old data"""
    try:
        data = request.get_json() or {}
        days_old = data.get('days_old', 30)
        
        orchestrator.cleanup_system(days_old)
        
        return jsonify({
            'message': f'System cleanup completed (removed data older than {days_old} days)',
            'status': 'success'
        })
    except Exception as e:
        logger.error(f"Error during cleanup: {str(e)}")
        return jsonify({'error': str(e)}), 500

# Serve React frontend in production (FRONTEND_BUILD_DIR set at app init)
@app.route('/', defaults={'path': ''})
@app.route('/<path:path>')
def serve_frontend(path):
    """Serve React frontend static files"""
    if path != "" and os.path.exists(os.path.join(FRONTEND_BUILD_DIR, path)):
        return send_from_directory(FRONTEND_BUILD_DIR, path)
    else:
        # Serve index.html for all non-API routes (React Router)
        if os.path.exists(os.path.join(FRONTEND_BUILD_DIR, 'index.html')):
            return send_from_directory(FRONTEND_BUILD_DIR, 'index.html')
        return jsonify({'error': 'Frontend not built. Run: cd frontend && npm run build'}), 404

if __name__ == '__main__':
    print("Starting Requirements Engineering API Server...")
    print(f"  Max request body size: {_max_mb} MB (override with MAX_CONTENT_MB env var)")
    print("API Endpoints:")
    print("  GET  /api/health - Health check")
    print("  POST /api/clarify-requirements - Clarify vague requirements")
    print("  POST /api/clarification-copilot - Live clarification assistant")
    print("  POST /api/clarification-copilot-turn - Conversational copilot turn")
    print("  POST /api/process-single - Process single requirement")
    print("  POST /api/process-and-generate-srs - Process text + generate SRS (single request)")
    print("  POST /api/process-audio - Process audio recording")
    print("  POST /api/transcribe-audio - Transcribe audio only (for live transcription)")
    print("  POST /api/process-batch - Process batch requirements")
    print("  POST /api/generate-srs - Generate SRS document")
    print("  POST /api/generate-srs-stream - Stream SRS generation (SSE)")
    print("  POST /api/expert-review/submit - Submit SRS for human expert review")
    print("  GET  /api/expert-review/requests - List expert review requests")
    print("  POST /api/expert-review/requests/<id>/messages - Append expert/user chat message")
    print("  PATCH /api/expert-review/requests/<id> - Complete expert review")
    print("  POST /api/evaluate-srs-kb-metrics - KB-style quality scores for raw SRS text (app use)")
    print("  POST /api/download-srs/<format> - Download SRS")
    print("  GET  /api/stats - Get system statistics")
    print("  POST /api/cleanup - Clean up system")
    
    # Check if frontend is built
    if os.path.exists(FRONTEND_BUILD_DIR):
        print(f"\nFrontend detected at: {FRONTEND_BUILD_DIR}")
        print("Serving frontend and API on http://localhost:8000")
    else:
        print(f"\nFrontend not found at: {FRONTEND_BUILD_DIR}")
        print("API only mode. Build frontend with: cd frontend && npm run build")
    
    port = int(os.environ.get('PORT', 8000))
    debug = os.environ.get('FLASK_ENV', 'production') != 'production'
    
    app.run(host='0.0.0.0', port=port, debug=debug)
